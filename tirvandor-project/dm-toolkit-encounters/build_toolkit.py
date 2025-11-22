#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("🔨 Building DM Toolkit - Random Encounters...")

doc = Document()

sections = doc.sections
for section in sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DM TOOLKIT\n\nRANDOM ENCOUNTERS")
run.font.size = Pt(36)
run.font.bold = True
doc.add_page_break()

filepath = 'markdown/01-encounter-tables.md'
if os.path.exists(filepath):
    print(f"  • 01-encounter-tables.md...", end='')
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Heading 1'
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
        elif line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Heading 3'
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            if p.runs:
                p.runs[0].font.size = Pt(11)
        elif line.strip():
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
        else:
            doc.add_paragraph()
    
    print(" ✅")

output_path = '/mnt/user-data/outputs/DM-Toolkit-Random-Encounters.docx'
doc.save(output_path)
size = os.path.getsize(output_path) / 1024
print(f"\n✅ COMPLETE: {size:.0f} KB")
