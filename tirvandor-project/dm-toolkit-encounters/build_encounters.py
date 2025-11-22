import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Use local markdown directory
script_dir = os.path.dirname(os.path.abspath(__file__))
markdown_dir = os.path.join(script_dir, "markdown")
cover_path = "/home/claude/tirvandor-project/tirvandor-cover-dm-toolkit-converted.png"
output_path = "/mnt/user-data/outputs/DM-Toolkit-Random-Encounters.docx"

def clean_text(text):
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('`', '').replace('#', '')
    text = text.replace('│', '').replace('┌', '').replace('┐', '')
    text = text.replace('└', '').replace('┘', '').replace('├', '').replace('┤', '')
    text = text.replace('─', '').strip()
    return text

def add_cover_and_copyright(doc):
    if os.path.exists(cover_path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(cover_path, width=Inches(8.5), height=Inches(11))
        doc.add_page_break()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    copyright_para = doc.add_paragraph()
    copyright_run = copyright_para.add_run("© 2024-2025 Tirvandor Campaign Setting. All rights reserved.\nFor personal tabletop use only. Not for commercial distribution.")
    copyright_run.font.size = Pt(8)
    copyright_run.italic = True
    copyright_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    copyright_para.paragraph_format.space_after = Pt(12)

def parse_table(lines, start_idx):
    idx = start_idx
    if idx >= len(lines) or '|' not in lines[idx]:
        return [], idx
    
    header_line = lines[idx].strip()
    headers = [h.strip() for h in header_line.split('|') if h.strip()]
    idx += 1
    
    if idx < len(lines) and '---' in lines[idx]:
        idx += 1
    
    rows = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not '|' in line or line.startswith('#'):
            break
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
        idx += 1
    
    return [headers] + rows if rows else [], idx

print("🔨 Building DM Toolkit - Random Encounters...")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

add_cover_and_copyright(doc)

doc.add_heading('Random Encounter Tables', 1).runs[0].font.color.rgb = RGBColor(139, 69, 19)
doc.add_paragraph("For Improvisation & Dynamic Gameplay")
doc.add_page_break()

with open(os.path.join(markdown_dir, "RANDOM-ENCOUNTER-TABLES.md"), 'r') as f:
    lines = f.readlines()

idx = 0
while idx < len(lines):
    line = lines[idx].strip()
    
    if not line:
        idx += 1
        continue
    
    if '|' in line and not line.startswith('#'):
        table_data, new_idx = parse_table(lines, idx)
        if table_data and len(table_data) > 1:
            num_cols = len(table_data[0])
            table = doc.add_table(rows=len(table_data), cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            for row_idx, row_data in enumerate(table_data):
                for col_idx in range(min(num_cols, len(row_data))):
                    table.rows[row_idx].cells[col_idx].text = row_data[col_idx]
                    for p in table.rows[row_idx].cells[col_idx].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
                            if row_idx == 0:
                                run.bold = True
            
            if num_cols == 2:
                table.rows[0].cells[0].width = Inches(0.8)
                table.rows[0].cells[1].width = Inches(5.7)
            
            idx = new_idx
        else:
            idx += 1
        continue
    
    if line.startswith('# '):
        doc.add_heading(clean_text(line[1:]), 1).runs[0].font.color.rgb = RGBColor(139, 69, 19)
    elif line.startswith('## '):
        doc.add_heading(clean_text(line[2:]), 2).runs[0].font.color.rgb = RGBColor(47, 79, 79)
    elif line.startswith('### '):
        doc.add_heading(clean_text(line[3:]), 3)
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(clean_text(line[2:]), style='List Bullet')
    elif not line.startswith('---'):
        doc.add_paragraph(clean_text(line))
    
    idx += 1

doc.save(output_path)
print(f"✅ Complete: {output_path}")
print(f"   Source: {markdown_dir}/RANDOM-ENCOUNTER-TABLES.md")
