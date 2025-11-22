#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os

print("🔨 Building Blood & Coin Players Guide...")

doc = Document()

sections = doc.sections
for section in sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add cover
print("  • Adding cover image...")
cover_path = 'images/cover.jpg'
if os.path.exists(cover_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(6.5))
    doc.add_page_break()
    print("    ✅ Cover added")

chapters = [
    '01-introduction.md', '02-contents.md', '03-campaign-overview.md',
    '04-character-creation.md', '05-reputation-system.md', '06-contract-types.md',
    '07-moral-choices.md', '08-campaign-structure.md', '09-starting-campaign.md',
    '10-character-hooks.md', '11-equipment.md', '12-house-rules.md', '13-legal.md'
]

base_dir = 'markdown'
first_chapter = True

for chapter_file in chapters:
    filepath = os.path.join(base_dir, chapter_file)
    if not os.path.exists(filepath):
        print(f"  ⚠️  Skipping {chapter_file}")
        continue
    
    print(f"  • {chapter_file}...", end='')
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table = None
    
    for line in lines:
        line = line.rstrip()
        
        # Detect table start
        if '|' in line and not in_table:
            # Start of table
            headers = [cell.strip() for cell in line.split('|')[1:-1]]
            in_table = True
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            continue
        
        # In table - skip separator line
        if in_table and '|---' in line or '|-' in line:
            continue
        
        # In table - add data row
        if in_table and '|' in line:
            cells_text = [cell.strip() for cell in line.split('|')[1:-1]]
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(cells_text):
                row_cells[i].text = cell_text
            continue
        
        # End of table
        if in_table and '|' not in line:
            in_table = False
            table = None
            doc.add_paragraph()  # Space after table
        
        if line.startswith('# '):
            if not first_chapter:
                doc.add_page_break()
            p = doc.add_paragraph(line[2:])
            p.style = 'Heading 1'
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            first_chapter = False
            
        elif line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Heading 3'
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            run = p.runs[0] if p.runs else None
            if run:
                run.font.size = Pt(11)
                
        elif line.strip():
            # Handle bold/italic markdown
            text = line
            p = doc.add_paragraph()
            
            # Split by ** for bold
            parts = text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(11)
                if i % 2 == 1:  # Odd indexes are bold
                    run.font.bold = True
            
            pf = p.paragraph_format
            pf.space_after = Pt(6)
        else:
            doc.add_paragraph()
    
    print(" ✅")

output_path = '/mnt/user-data/outputs/Blood-and-Coin-Players-Guide-Rebuilt.docx'
doc.save(output_path)

size = os.path.getsize(output_path)
size_mb = size / (1024 * 1024)
print(f"\n✅ COMPLETE: {size_mb:.2f} MB")
