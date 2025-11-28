import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

world_dir = "markdown"
images_dir = "images/world-images"
cover_path = "images/world-guide-cover-converted.png"

seen_headers = set()
seen_images = set()

def clean(text):
    return re.sub(r'[^\x00-\x7F]+', '', text).replace('**', '').replace('`', '').replace('#', '').strip()

def parse_table(lines, idx):
    table_lines = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        if '|---' not in line:
            table_lines.append(line)
        idx += 1
    if len(table_lines) < 2:
        return None, idx
    data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    max_cols = max(len(row) for row in data)
    for row in data:
        while len(row) < max_cols:
            row.append('')
    return data, idx

def parse_roll_table(lines, idx):
    """Parse numbered lists after 'Roll d20' or similar"""
    table_data = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or line.startswith('#') or line.startswith('**'):
            break
        match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if match:
            roll = match.group(1)
            detail = match.group(2)
            table_data.append([roll, detail])
            idx += 1
        else:
            break
    return table_data if table_data else None, idx

def add_image(doc, filename, width=5.5):
    if filename in seen_images:
        return False
    # Try world-images first, then world-maps
    path = os.path.join(images_dir, filename)
    if not os.path.exists(path):
        path = os.path.join("images/world-maps", filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        p.paragraph_format.space_after = Pt(12)
        seen_images.add(filename)
        return True
    return False

def match_image(header_text):
    """Match header to image by keywords"""
    h = header_text.lower()
    
    # Exact matches from filename to likely header text
    mapping = {
        'geography': 'tirvandor-world-geography-chapter-opener.jpg',
        'pantheon': 'tirvandor-world-pantheon-chapter-opener.jpg',
        'timeline': 'tirvandor-world-history-timeline-chapter-opener.jpg',
        'history': 'tirvandor-world-history-timeline-chapter-opener.jpg',
        'factions': 'tirvandor-world-factions-chapter-opener.jpg',
        'noble houses': 'tirvandor-world-noble-houses-chapter-opener.jpg',
        'houses and dynasties': 'tirvandor-world-noble-houses-chapter-opener.jpg',
        
        'aethermere (capital': 'tirvandor-world-aethermere.jpg',
        'goldreach': 'tirvandor-world-goldreach.jpg',
        'kaer thandros': 'tirvandor-world-kaer-thandros.jpg',
        'ironhold': 'tirvandor-world-ironhold-dwarven-mountain-city.jpg',
        'silverpine': 'tirvandor-world-silverpine-university-city.jpg',
        
        'aethoria (eastern continent)': 'tirvandor-world-aethorias-free-lands.jpg',
        'sundering sea': 'tirvandor-world-the-sundering-sea.jpg',
        'shattered shore': 'tirvandor-world-the-shattered-coast.jpg',
        'shattered coast': 'tirvandor-world-the-shattered-coast.jpg',
        'contested lands': 'tirvandor-world-the-contested-lands-no-mans-land.jpg',
        'ironspine holds': 'tirvandor-world-ironhold-mountain-range.jpg',
        'silverwood': 'tirvandor-world-silverpine-ancient-forest.jpg',
        'thaldros (western': 'tirvandor-world-thaldros-military-borderlands.jpg',
        'frostmarches': 'tirvandor-world-thaldros-military-borderlands.jpg',
        
        'the sundering': 'tirvandor-world-the-sundering.jpg',
        'pre-sundering': 'tirvandor-world-pre-sundering-unity.jpg',
        'great accord': 'tirvandor-world-great-accord-signing.jpg',
        'healing of the rift': 'tirvandor-world-healing-of-the-rift.jpg',
        'blackwood': 'tirvandor-world-house-blackwoods-execution.jpg',
        
        'the seven': 'tirvandor-world-seven-gods-council.jpg',
        'the prime deities': 'tirvandor-world-seven-gods-council.jpg',
        'prime deities': 'tirvandor-world-seven-gods-council.jpg',
        'divine magic in practice': 'tirvandor-world-ley-line-nexus.jpg',
        
        'the iron guild': 'tirvandor-world-iron-guild-emblem.jpg',
        'iron guild': 'tirvandor-world-iron-guild-emblem.jpg',
        'guildhalls': 'tirvandor-world-iron-guild-hall-interior.jpg',
        'guild headquarters': 'tirvandor-world-iron-guild-hall-interior.jpg',
        'the silver circle': 'tirvandor-world-silver-circle-emblem.jpg',
        'silver circle': 'tirvandor-world-silver-circle-emblem.jpg',
        'thornwood syndicate': 'tirvandor-world-thornwood-syndicate-emblem.jpg',
        'the thornwood': 'tirvandor-world-thornwood-syndicate-emblem.jpg',
        
        'natural magic': 'tirvandor-world-spellcasting-with-ley-lines.jpg',
        'magical societies': 'tirvandor-world-magic-and-ley-lines-chapter-opener.jpg',
        'known ley line nexuses': '../world-maps/tirvandor_map_ley_lines.jpg',
        'ley line nexus': 'tirvandor-world-ley-line-convergence-detailed.jpg',
        'major roads of thaldros': '../world-maps/tirvandor_map_locations_thaldros.jpg',
        'major roads of aethoria': '../world-maps/tirvandor_map_locations_aethoria.jpg',
        'major roads': '../world-maps/tirvandor_map_locations.jpg',
        'ley line network': '../world-maps/tirvandor_map_ley_lines.jpg',
        'regions of tirvandor': '../world-maps/tirvandor_map_regions.jpg',
        'sundering sea': '../world-maps/tirvandor_map_locations_sundering_sea.jpg',
        'ley line convergence': 'tirvandor-world-ley-line-convergence-detailed.jpg',
        
        'ancient primordials': 'tirvandor-world-ancient-precursor-ruins.jpg',
        'the primordials': 'tirvandor-world-ancient-precursor-ruins.jpg',
        'precursor architecture': 'tirvandor-world-ancient-precursor-architecture-detail.jpg',
        'ancient ruins': 'tirvandor-world-ancient-precursor-architecture-detail.jpg',
        'precursors to the sundering': 'tirvandor-world-ancient-precursor-ruins.jpg',
        
        'border village': 'tirvandor-world-border-town-contested-lands.jpg',
        'border settlements': 'tirvandor-world-border-town-contested-lands.jpg',
        'refugee camps': 'tirvandor-world-refugee-camp-scene.jpg',
        'refugee crisis': 'tirvandor-world-refugee-camp-scene.jpg',
        'void rift': 'tirvandor-world-void-rift-manifestation.jpg',
        'the void': 'tirvandor-world-void-rift-manifestation.jpg',
        'border wraith': 'tirvandor-world-border-wraith.jpg',
        'wraiths': 'tirvandor-world-border-wraith.jpg',
        'sundering scar': 'tirvandor-world-sundering-scar.jpg',
        
        'throne room': 'tirvandor-world-kaer-thandros-throne-room.jpg',
        'royal palace': 'tirvandor-world-kaer-thandros-throne-room.jpg',
        'council grove': 'tirvandor-world-goldreach-city-council.jpg',
        'city council': 'tirvandor-world-goldreach-city-council.jpg',
        'the docks': 'tirvandor-world-goldreach-docks-at-dawn.jpg',
        'harbor district': 'tirvandor-world-goldreach-docks-at-dawn.jpg',
        'skyline': 'tirvandor-world-goldreach-skyline.jpg',
        'goldcoast': 'tirvandor-world-goldreach-skyline.jpg',
        'heartlands': 'tirvandor-world-goldreach-docks-at-dawn.jpg',
    }
    
    for key, img in mapping.items():
        if key in h:
            return img
    return None

def is_dup(header_text, level):
    # Don't dedupe contextual headers that appear in multiple sections
    exempt = ['adventure hooks', 'description', 'key locations', 'notable npcs', 'rumors', 'secrets', 
              'current status', 'conflicts & alliances', 'atmosphere', 'culture & people', 'foundation',
              'approaching from', 'during session', 'the settlement']
    if any(ex in header_text.lower() for ex in exempt):
        return False
    
    key = f"{level}:{header_text.lower().strip()}"
    if key in seen_headers:
        return True
    seen_headers.add(key)
    return False

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

if os.path.exists(cover_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(8.5), height=Inches(11))
    doc.add_page_break()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

cp = doc.add_paragraph()
cr = cp.add_run("© 2024-2025 Tirvandor Campaign Setting.\nFor personal tabletop use only.")
cr.font.size = Pt(8)
cr.italic = True
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('TIRVANDOR WORLD GUIDE', 1).runs[0].font.color.rgb = RGBColor(139, 69, 19)
doc.add_paragraph("The Realm of Deep Magic")
doc.add_page_break()

toc = doc.add_heading('TABLE OF CONTENTS', 2)
toc.runs[0].font.color.rgb = RGBColor(47, 79, 79)
for i, t in enumerate(['Geography', 'Pantheon', 'Timeline', 'Factions', 'Houses', 'Settlements', 'Settlement Histories', 'Locations', 'Street Details', 'Lore'], 1):
    doc.add_paragraph(f"{i}. {t}")
doc.add_page_break()

files = [
    '01-introduction.md',
    '02-geography.md', 
    '03-timeline.md',
    '04-pantheon.md',
    '05-factions.md',
    '06-noble-houses.md',
    '07-settlement-registry.md',
    '08-settlement-histories.md',
    '09-locations.md',
    '10-regional-lore.md',
    '11-street-level-details.md'
]

for filename in files:
    filepath = os.path.join(world_dir, filename)
    if not os.path.exists(filepath):
        continue
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line or line.startswith('---'):
            idx += 1
            continue
        
        if line.startswith('|'):
            table_data, idx = parse_table(lines, idx)
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                for r, row_data in enumerate(table_data):
                    for c, cell_data in enumerate(row_data):
                        table.rows[r].cells[c].text = clean(cell_data)
                        for p in table.rows[r].cells[c].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                                if r == 0:
                                    run.bold = True
            continue
        
        # Random Detail Tables or Roll d20 instructions
        if 'roll d20' in line.lower() or 'random detail table' in line.lower():
            p = doc.add_paragraph(clean(line))
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.bold = True
            idx += 1
            
            # Parse following numbered list as roll table
            roll_data, idx = parse_roll_table(lines, idx)
            if roll_data:
                table = doc.add_table(rows=len(roll_data)+1, cols=2)
                table.style = 'Light Grid Accent 1'
                table.rows[0].cells[0].text = 'd20'
                table.rows[0].cells[1].text = 'Detail'
                for r_idx, (roll, detail) in enumerate(roll_data):
                    table.rows[r_idx+1].cells[0].text = roll
                    table.rows[r_idx+1].cells[1].text = clean(detail)
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                    if row == table.rows[0]:
                        for p in row.cells[0].paragraphs:
                            for run in p.runs:
                                run.bold = True
                        for p in row.cells[1].paragraphs:
                            for run in p.runs:
                                run.bold = True
                table.rows[0].cells[0].width = Inches(0.6)
                table.rows[0].cells[1].width = Inches(5.9)
            continue
        
        if line.startswith('# '):
            if is_dup(line[1:], 1):
                idx += 1
                continue
            h = doc.add_heading(clean(line[1:]), 1)
            h.runs[0].font.color.rgb = RGBColor(139, 69, 19)
            img = match_image(line)
            if img:
                add_image(doc, img, 6.0)
        elif line.startswith('## '):
            if is_dup(line[2:], 2):
                idx += 1
                continue
            h = doc.add_heading(clean(line[2:]), 2)
            h.runs[0].font.color.rgb = RGBColor(47, 79, 79)
            img = match_image(line)
            if img:
                add_image(doc, img, 5.5)
        elif line.startswith('### '):
            if is_dup(line[3:], 3):
                idx += 1
                continue
            h = doc.add_heading(clean(line[3:]), 3)
            img = match_image(line)
            if img:
                add_image(doc, img, 5.0)
        elif line.startswith('#### '):
            if is_dup(line[4:], 4):
                idx += 1
                continue
            doc.add_heading(clean(line[4:]), 4)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(clean(line[2:]), style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
        elif re.match(r'^\d+\.\s', line):
            # Check if it's a sequential list that should stay numbered
            if any(x in line.lower() for x in ['step', 'first', 'second', 'third', 'then', 'finally']):
                # Keep as numbered list
                p = doc.add_paragraph(clean(line[line.index('.')+1:].strip()), style='List Number')
            else:
                # Convert to plain text with number removed
                p = doc.add_paragraph(clean(line))
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph(clean(line))
            p.paragraph_format.space_after = Pt(2)
        
        idx += 1
    
    doc.add_page_break()

doc.save("/mnt/user-data/outputs/Tirvandor-World-Guide.docx")
print(f"✅ {len(seen_images)} images placed contextually")
