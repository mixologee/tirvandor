import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Use local markdown directory
script_dir = os.path.dirname(os.path.abspath(__file__))
markdown_dir = os.path.join(script_dir, "markdown")
npc_images_dir = "/home/claude/tirvandor-project/npc-portraits"
cover_path = "/home/claude/tirvandor-project/tirvandor-cover-dm-toolkit-converted.png"
output_path = "/mnt/user-data/outputs/DM-Toolkit-NPC-Cards.docx"

def clean_text(text):
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('`', '').replace('#', '')
    text = text.replace('│', '').replace('┌', '').replace('┐', '').replace('└', '').replace('┘', '').replace('├', '').replace('┤', '').replace('─', '').strip()
    return text

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)

def find_image(filename):
    if not filename: return None
    path = os.path.join(npc_images_dir, filename)
    if os.path.exists(path): return path
    base = os.path.splitext(filename)[0]
    for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
        path = os.path.join(npc_images_dir, base + ext)
        if os.path.exists(path): return path
    return None

def add_cover_and_copyright(doc):
    if os.path.exists(cover_path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(cover_path, width=Inches(8.5), height=Inches(11))
        doc.add_page_break()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    copyright_para = doc.add_paragraph()
    copyright_run = copyright_para.add_run("© 2024-2025 Tirvandor Campaign Setting. All rights reserved.\nFor personal tabletop use only. Not for commercial distribution.")
    copyright_run.font.size = Pt(8)
    copyright_run.italic = True
    copyright_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

print("🔨 Building DM Toolkit - NPC Quick Cards...")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

add_cover_and_copyright(doc)
doc.add_heading('NPC Quick Reference Cards', 1).runs[0].font.color.rgb = RGBColor(139, 69, 19)
doc.add_paragraph("Essential NPCs for Tirvandor Campaigns - Top 20")
doc.add_page_break()

with open(os.path.join(markdown_dir, "NPC-QUICK-CARDS-TOP-20.md"), 'r') as f:
    lines = f.readlines()

cards, current_card_lines, current_portrait, in_card = [], [], None, False
skip_first_h2 = True

for line in lines:
    line_stripped = line.strip()
    if line_stripped.startswith('## ') and not line_stripped.startswith('###'):
        if skip_first_h2:
            skip_first_h2 = False
            continue
        if current_card_lines:
            cards.append({'header': current_card_lines[0], 'portrait': current_portrait, 'content': current_card_lines[1:]})
        current_card_lines, current_portrait, in_card = [line_stripped], None, True
    elif in_card:
        if '**Portrait:**' in line_stripped:
            current_portrait = line_stripped.split('`')[1] if '`' in line_stripped else None
        elif not line_stripped.startswith('---') and not line_stripped.startswith('**Format') and not line_stripped.startswith('**Print'):
            if line_stripped: current_card_lines.append(line_stripped)

if current_card_lines:
    cards.append({'header': current_card_lines[0], 'portrait': current_portrait, 'content': current_card_lines[1:]})

print(f"Found {len(cards)} NPC cards")

images_added, first_card = 0, True

for card in cards:
    if not first_card: doc.add_page_break()
    first_card = False
    
    header_text = clean_text(card['header'][3:]) if card['header'].startswith('## ') else clean_text(card['header'])
    doc.add_heading(header_text, 2).runs[0].font.color.rgb = RGBColor(47, 79, 79)
    
    table = doc.add_table(rows=1, cols=2)
    table.autofit, table.allow_autofit = False, False
    
    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(3.2)
    
    if card['portrait']:
        image_path = find_image(card['portrait'])
        if image_path:
            try:
                p = left_cell.paragraphs[0]
                run = p.add_run()
                run.add_picture(image_path, width=Inches(3.0), height=Inches(4.5))
                images_added += 1
            except: left_cell.text = f"[Portrait: {card['portrait']}]"
        else: left_cell.text = f"[Portrait: {card['portrait']}]"
    
    right_cell = table.rows[0].cells[1]
    right_cell.width = Inches(3.8)
    right_cell.paragraphs[0].text = ''
    
    for content_line in card['content']:
        if content_line.startswith('- ') or content_line.startswith('* '):
            p = right_cell.add_paragraph('', style='List Bullet')
            add_formatted_text(p, content_line[2:])
        else:
            p = right_cell.add_paragraph()
            add_formatted_text(p, content_line)
        p.paragraph_format.space_after = Pt(2)
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find('.//w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)

doc.save(output_path)
print(f"✅ Complete: {output_path}")
print(f"   Source: {markdown_dir}/NPC-QUICK-CARDS-TOP-20.md")
print(f"   Images: {images_added}/{len(cards)}")
