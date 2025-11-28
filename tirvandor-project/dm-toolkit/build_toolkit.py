#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("🔨 Building DM Toolkit - Reference Tools...")

doc = Document()

sections = doc.sections
for section in sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DM TOOLKIT\n\nREFERENCE TOOLS")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)
doc.add_page_break()

chapters = [
    '01-dm-quick-reference.md',
    '02-location-quick-reference.md',
    '03-faction-relationship-map.md',
    '04-monster-stats-condensed.md',
    '05-skill-challenge-templates.md',
    '06-session-tracking.md',
    '07-timeline-visual.md',
    '08-campaign-dashboard.md',
    '09-session-zero.md',
    '10-atmosphere-guide.md'
]

base_dir = 'markdown'
first_chapter = True

for chapter_file in chapters:
    filepath = os.path.join(base_dir, chapter_file)
    if not os.path.exists(filepath):
        print(f"  ⚠️  Skipping {chapter_file}")
        continue
    
    print(f"  • {chapter_file}...", end='')
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            if not first_chapter:
                doc.add_page_break()
            p = doc.add_paragraph(line[2:])
            p.style = 'Heading 1'
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
            first_chapter = False
            
        elif line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True
            
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Heading 3'
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            if p.runs:
                p.runs[0].font.size = Pt(11)
                
        elif line.strip():
            p = doc.add_paragraph(line)
            pf = p.paragraph_format
            pf.space_after = Pt(6)
            if p.runs:
                p.runs[0].font.size = Pt(11)
        else:
            doc.add_paragraph()
    
    print(" ✅")

output_path = '/mnt/user-data/outputs/DM-Toolkit-Reference-Tools.docx'
doc.save(output_path)

size = os.path.getsize(output_path)
size_kb = size / 1024
print(f"\n✅ COMPLETE: {size_kb:.0f} KB")
