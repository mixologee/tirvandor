import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Cover
for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover = 'images/tirvandor-cover-blood-and-coin-campaign.png'
if os.path.exists(cover):
    p.add_run().add_picture(cover, width=Inches(8.5))
doc.add_page_break()

# Reset margins
for s in doc.sections: s.top_margin = s.bottom_margin = Inches(0.5); s.left_margin = s.right_margin = Inches(0.75)

def clean(t): return re.sub(r'[^\x00-\x7F]+','',t).replace('**','').replace('`','').strip()

def add_img(path, w=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
        return True
    return False

# Process Act I
for f in ['ACT-I-COMPLETE.md'] + [f'QUEST-2.{i}-{n}.md' for i,n in [(1,'RECRUITERS'),(2,'FIRST-MISSION'),(3,'DOUBLE-CROSS'),(4,'THE-HEIST'),(5,'BREAKING-POINT')]] + [f'QUEST-3.{i}-{n}.md' for i,n in [(1,'PEACEKEEPERS'),(2,'UNDERGROUND'),(3,'ALLIANCE'),(4,'UNIFICATION'),(5,'MEDIATORS-REST')]]:
    fpath = f'markdown/{f}'
    if not os.path.exists(fpath): continue
    with open(fpath, 'r', encoding='utf-8', errors='ignore') as file:
        for line in file:
            line = line.rstrip()
            if not line or line.startswith('---'): continue
            if line.startswith('# '): doc.add_heading(clean(line[2:]), 1).runs[0].font.color.rgb = RGBColor(139,0,0)
            elif line.startswith('## '): doc.add_heading(clean(line[3:]), 2).runs[0].font.color.rgb = RGBColor(47,79,79)
            elif line.startswith('### '): doc.add_heading(clean(line[4:]), 3)
            elif line.startswith('- '): doc.add_paragraph(clean(line[2:]), style='List Bullet')
            elif clean(line): doc.add_paragraph(clean(line))

# Add images from directories
for subdir in ['npcs', 'items', 'maps']:
    img_dir = f'images/{subdir}'
    if os.path.exists(img_dir):
        for img in sorted(os.listdir(img_dir))[:20]:
            if img.endswith(('.jpg','.png')):
                add_img(f'{img_dir}/{img}')

doc.save('/mnt/user-data/outputs/Blood-and-Coin-Campaign.docx')
print("✅ Blood & Coin rebuilt")
