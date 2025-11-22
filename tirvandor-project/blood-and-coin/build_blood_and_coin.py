import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

bc_dir = "markdown"
images_dir = "images"
cover_path = "images/blood-and-coin-cover.jpg"

seen_headers = set()
seen_images = set()

def clean(text):
    # Remove non-ASCII, then remove markdown formatting
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('**', '').replace('*', '').replace('`', '').replace('#', '')
    return text.strip()

def parse_table(lines, idx):
    """Parse markdown table"""
    table_lines = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        if '|---' not in line:
            table_lines.append(line)
        idx += 1
    if len(table_lines) < 2:
        return None, idx
    data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    max_cols = max(len(row) for row in data) if data else 0
    for row in data:
        while len(row) < max_cols:
            row.append('')
    return data, idx

def parse_table(lines, idx):
    """Parse markdown table"""
    table_lines = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        if '|---' not in line and '|-' not in line:  # Skip separator lines
            table_lines.append(line)
        idx += 1
    if len(table_lines) < 2:
        return None, idx
    data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    max_cols = max(len(row) for row in data) if data else 0
    for row in data:
        while len(row) < max_cols:
            row.append('')
    return data, idx

def add_image(doc, filename, width=5.5):
    if filename in seen_images:
        return False
    # Try multiple locations
    paths = [
        os.path.join(images_dir, filename),
        os.path.join(images_dir, 'npcs', filename),
        os.path.join(images_dir, 'maps', filename),
        os.path.join(images_dir, 'items', filename)
    ]
    for path in paths:
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
            p.paragraph_format.space_after = Pt(12)
            seen_images.add(filename)
            return True
    return False

def match_image(header_text, quest_context=""):
    """Match headers and context to appropriate images"""
    h = header_text.lower()
    q = quest_context.lower()
    
    # NPC mappings - comprehensive list
    npc_mapping = {
        'elise consortium': 'bc-elise-thornwood-moonlight.jpg',
        'elise thornwood': 'bc-elise-thornwood-reaper.jpg',
        'catherine': 'faction-lyanna-thornwood-red-wolf.jpg',
        'red wolf': 'faction-lyanna-thornwood-red-wolf.jpg',
        'corvus blackwood': 'bc-corvus-blackwood-necromancer.jpg',
        'corvus': 'bc-corvus-blackwood-necromancer.jpg',
        'lord shadows': 'bc-corvus-blackwood-necromancer.jpg',
        'kael shadowstep': 'bc-kael-shadowstep-rogue.jpg',
        'kael shadowbane': 'bc-kael-shadowstep-rogue.jpg',
        'kael': 'bc-kael-shadowstep-rogue.jpg',
        'lyanna thornwood': 'faction-lyanna-thornwood-red-wolf.jpg',
        'lyanna': 'faction-lyanna-thornwood-red-wolf.jpg',
        'helena dawnblade': 'faction-helena-blackstone-iron-legion.jpg',
        'dawnblade': 'faction-helena-blackstone-iron-legion.jpg',
        'commander dawnblade': 'faction-helena-blackstone-iron-legion.jpg',
        'commander helena': 'faction-helena-blackstone-iron-legion.jpg',
        'aria dawnbringer': 'faction-aria-dawnbringer-paladin.jpg',
        'aria': 'faction-aria-dawnbringer-paladin.jpg',
        'garrick ironheart': 'faction-garrick-ironheart-guildmaster.jpg',
        'garrick': 'faction-garrick-ironheart-guildmaster.jpg',
        'roderic ironfist': 'faction-roderic-ironfist-iron-guild.jpg',
        'roderic': 'faction-roderic-ironfist-iron-guild.jpg',
        'varak ironfist': 'faction-roderic-ironfist-iron-guild.jpg',
        'varak': 'faction-roderic-ironfist-iron-guild.jpg',
        'elara': 'bc-elara-prophet.jpg',
        'elara corvus': 'bc-elara-prophet.jpg',
        'arcanus': 'bc-arcanus-elder-lich.jpg',
        'faceless': 'bc-faceless-assassin.jpg',
        'faceless assassin': 'bc-faceless-assassin.jpg',
        'viktor seaworth': 'gr-admiral-viktor-seaworth.jpg',
        'admiral viktor': 'gr-admiral-viktor-seaworth.jpg',
    }
    
    # Map mappings
    map_mapping = {
        'broken spear': 'broken-spear-tavern.jpg',
        'tavern encounter': 'broken-spear-tavern.jpg',
        'safehouse': 'criminal-safehouse.jpg',
        'smuggler': 'dockside-smuggler-den.jpg',
        'warehouse': 'warehouse-district-ambush-night.jpg',
        'ambush': 'warehouse-district-ambush-night.jpg',
        'sewers': 'goldreach-sewers.jpg',
        'iron guild hall': 'iron-guild-hall.jpg',
        'guild hall': 'iron-guild-hall.jpg',
        'final showdown': 'final-showdown-arena.jpg',
        'final battle': 'final-showdown-arena.jpg',
        'arena': 'final-showdown-arena.jpg',
    }
    
    # Check NPCs first
    for key, img in npc_mapping.items():
        if key in h:
            return img
    
    # Check maps
    for key, img in map_mapping.items():
        if key in h or key in q:
            return img
    
    return None

def is_dup(header_text, level):
    exempt = ['description', 'tactics', 'rewards', 'consequences', 'read aloud', 
              'setup', 'encounter', 'conclusion', 'background', 'notes']
    if any(ex in header_text.lower() for ex in exempt):
        return False
    
    key = f"{level}:{header_text.lower().strip()}"
    if key in seen_headers:
        return True
    seen_headers.add(key)
    return False

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

# Cover page
if os.path.exists(cover_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(8.5))
    doc.add_page_break()

# Set normal margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Copyright
cp = doc.add_paragraph()
cr = cp.add_run("© 2024-2025 Tirvandor Campaign Setting.\nFor personal tabletop use only.")
cr.font.size = Pt(8)
cr.italic = True
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Title
doc.add_heading('BLOOD & COIN CAMPAIGN', 1).runs[0].font.color.rgb = RGBColor(139, 0, 0)
doc.add_paragraph("A Morally Grey Campaign for Levels 1-15")
doc.add_page_break()

# TOC
toc = doc.add_heading('TABLE OF CONTENTS', 2)
toc.runs[0].font.color.rgb = RGBColor(47, 79, 79)
for i, t in enumerate(['Campaign Overview', 'Act I Quests', 'Act II Quests', 'Act III Quests', 
                       'NPCs', 'Monsters', 'Items', 'Handouts'], 1):
    doc.add_paragraph(f"{i}. {t}")
doc.add_page_break()

# File order
files = [
    'ACT-I-COMPLETE.md',
    'QUEST-2.1-RECRUITERS.md',
    'QUEST-2.2-FIRST-MISSION.md',
    'QUEST-2.3-DOUBLE-CROSS.md',
    'QUEST-2.4-THE-HEIST.md',
    'QUEST-2.5-BREAKING-POINT.md',
    'QUEST-3.1-PEACEKEEPERS.md',
    'QUEST-3.2-UNDERGROUND.md',
    'QUEST-3.3-ALLIANCE.md',
    'QUEST-3.4-UNIFICATION.md',
    'QUEST-3.5-MEDIATORS-REST.md',
    'npcs.md',
    'monsters.md',
    'items.md',
    'handouts.md'
]

current_quest = ""

for filename in files:
    filepath = os.path.join(bc_dir, filename)
    if not os.path.exists(filepath):
        continue
    
    if filename.startswith('QUEST'):
        current_quest = filename
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line or line.startswith('---'):
            idx += 1
            continue
        
        # Check for tables
        if line.startswith('|'):
            table_data, idx = parse_table(lines, idx)
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                for r, row_data in enumerate(table_data):
                    for c, cell_data in enumerate(row_data):
                        table.rows[r].cells[c].text = clean(cell_data)
                        for p in table.rows[r].cells[c].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                                if r == 0:
                                    run.bold = True
            continue
        
        if line.startswith('# '):
            if is_dup(line[1:], 1):
                idx += 1
                continue
            h = doc.add_heading(clean(line[1:]), 1)
            h.runs[0].font.color.rgb = RGBColor(139, 0, 0)
            img = match_image(line, current_quest)
            if img:
                add_image(doc, img, 6.0)
        elif line.startswith('## '):
            if is_dup(line[2:], 2):
                idx += 1
                continue
            h = doc.add_heading(clean(line[2:]), 2)
            h.runs[0].font.color.rgb = RGBColor(47, 79, 79)
            img = match_image(line, current_quest)
            if img:
                add_image(doc, img, 5.5)
        elif line.startswith('### '):
            if is_dup(line[3:], 3):
                idx += 1
                continue
            h = doc.add_heading(clean(line[3:]), 3)
            img = match_image(line, current_quest)
            if img:
                add_image(doc, img, 5.0)
        elif line.startswith('#### '):
            if is_dup(line[4:], 4):
                idx += 1
                continue
            doc.add_heading(clean(line[4:]), 4)
        elif line.startswith('> '):
            # Read-aloud text - strip the > and format
            text = clean(line[2:])
            if text:  # Only add non-empty lines
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(70, 70, 70)
        elif line.strip() == '>':
            # Empty blockquote line - skip it
            idx += 1
            continue
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(clean(line[2:]), style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph(clean(line))
            p.paragraph_format.space_after = Pt(2)
        
        idx += 1
    
    doc.add_page_break()

doc.save("/mnt/user-data/outputs/Blood-and-Coin-Campaign.docx")
print(f"✅ {len(seen_images)} images placed")
