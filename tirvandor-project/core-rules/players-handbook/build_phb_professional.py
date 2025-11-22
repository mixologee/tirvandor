#!/usr/bin/env python3
"""
Professional Player's Handbook Builder
Builds publication-quality D&D Player's Handbook from organized chapter files
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class PHBBuilder:
    def __init__(self):
        self.base_dir = "/home/claude/tirvandor-project/core-rules/players-handbook"
        self.markdown_dir = f"{self.base_dir}/markdown"
        self.images_dir = f"{self.base_dir}/images"
        self.output_path = "/mnt/user-data/outputs/Tirvandor-Players-Handbook-v2.docx"
        
        self.chapters = [
            '01-introduction.md',
            '02-character-creation.md',
            '03-races.md',
            '04-classes.md',
            '05-backgrounds.md',
            '06-equipment.md',
            '07-customization.md',
            '08-custom-subclasses.md',
            '09-custom-spells.md',
            '10-conditions.md',
            '11-world-primer.md',
            '12-legal.md'
        ]
        
    def create_document(self):
        """Create new document with proper formatting"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        return doc
    
    def add_cover_page(self, doc):
        """Add professional cover page"""
        print("  Adding cover page...")
        
        # Try to add cover image
        cover_path = f"{self.images_dir}/tirvandor-cover-players-guide.png"
        if os.path.exists(cover_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(cover_path, width=Inches(6.0))
            except Exception as e:
                print(f"    ⚠️  Cover image error: {e}")
                # Fallback to text cover
                self.add_text_cover(doc)
        else:
            self.add_text_cover(doc)
        
        doc.add_page_break()
    
    def add_text_cover(self, doc):
        """Add text-based cover if image fails"""
        # Title
        title = doc.add_heading("TIRVANDOR", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(48)
        title.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        
        # Subtitle
        subtitle = doc.add_heading("Player's Handbook", 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(32)
        subtitle.runs[0].font.color.rgb = RGBColor(47, 79, 79)
        
        # Tagline
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("Core Rules for Character Creation and Play")
        run.font.size = Pt(16)
        run.italic = True
    
    def process_markdown_line(self, line, doc):
        """Process a single markdown line and add to document"""
        line = line.rstrip()
        
        # Skip empty lines and separators
        if not line or line.strip() == '---':
            return
        
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), 1)
            heading.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), 2)
            heading.runs[0].font.color.rgb = RGBColor(47, 79, 79)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), 3)
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Tables (simple detection)
        elif line.startswith('|'):
            # Skip table separator lines
            if not re.match(r'^\|[\s\-:]+\|', line):
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                # Store for table processing (simplified for now)
                para = doc.add_paragraph(' | '.join(cells))
                para.paragraph_format.left_indent = Inches(0.25)
        
        # Regular paragraphs
        elif line.strip():
            # Clean markdown
            text = line.replace('**', '').replace('*', '').replace('`', '')
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
    
    def add_chapter(self, doc, chapter_file):
        """Add a chapter from markdown file"""
        filepath = f"{self.markdown_dir}/{chapter_file}"
        chapter_name = chapter_file.replace('.md', '').replace('-', ' ').title()
        print(f"  Adding: {chapter_name}")
        
        if not os.path.exists(filepath):
            print(f"    ⚠️  File not found: {chapter_file}")
            return
        
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    self.process_markdown_line(line, doc)
        except Exception as e:
            print(f"    ⚠️  Error processing {chapter_file}: {e}")
    
    def build(self):
        """Main build process"""
        print("=" * 60)
        print("BUILDING TIRVANDOR PLAYER'S HANDBOOK")
        print("=" * 60)
        print()
        
        # Create document
        print("📘 Creating document...")
        doc = self.create_document()
        
        # Add cover
        self.add_cover_page(doc)
        
        # Add all chapters
        print("📚 Adding chapters...")
        for chapter in self.chapters:
            self.add_chapter(doc, chapter)
        
        # Save
        print()
        print("💾 Saving document...")
        doc.save(self.output_path)
        
        # Report
        size_mb = os.path.getsize(self.output_path) / (1024 * 1024)
        print()
        print("=" * 60)
        print("✅ BUILD COMPLETE!")
        print("=" * 60)
        print(f"Output: {self.output_path}")
        print(f"Size: {size_mb:.2f} MB")
        print()
        print("📖 Player's Handbook ready for use!")

if __name__ == "__main__":
    builder = PHBBuilder()
    builder.build()
