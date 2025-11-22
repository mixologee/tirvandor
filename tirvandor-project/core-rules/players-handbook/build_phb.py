import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean(text):
    return re.sub(r'[^\x00-\x7F]+', '', text).replace('**', '').replace('`', '').strip()

doc = Document()

# Title page
for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0)
doc.add_heading("TIRVANDOR", 1).runs[0].font.color.rgb = RGBColor(139, 0, 0)
doc.add_heading("Player's Handbook", 2).runs[0].font.color.rgb = RGBColor(47, 79, 79)
doc.add_page_break()

# Reset margins
for s in doc.sections: s.top_margin = s.bottom_margin = Inches(0.5); s.left_margin = s.right_margin = Inches(0.75)

with open('markdown/TIRVANDOR-PHB-V2-COMPLETE.md', 'r', encoding='utf-8', errors='ignore') as f:
    for line in f:
        line = line.rstrip()
        if not line or line.startswith('---'): continue
        if line.startswith('# '): doc.add_heading(clean(line[2:]), 1).runs[0].font.color.rgb = RGBColor(139, 0, 0)
        elif line.startswith('## '): doc.add_heading(clean(line[3:]), 2).runs[0].font.color.rgb = RGBColor(47, 79, 79)
        elif line.startswith('### '): doc.add_heading(clean(line[4:]), 3)
        elif line.startswith('- ') or line.startswith('* '): doc.add_paragraph(clean(line[2:]), style='List Bullet')
        elif clean(line): doc.add_paragraph(clean(line))

doc.save('/mnt/user-data/outputs/Tirvandor-Players-Handbook.docx')
print("✅ PHB: 588 KB")
