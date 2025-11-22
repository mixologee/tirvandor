#!/usr/bin/env python3
"""
Complete PHB Builder with FULL Table of Contents
Shows all major sections
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class FullTOCBuilder:
    def __init__(self):
        self.base_dir = "/home/claude/tirvandor-project/core-rules/players-handbook"
        self.markdown_dir = f"{self.base_dir}/markdown"
        self.images_dir = f"{self.base_dir}/images"
        self.output_path = "/mnt/user-data/outputs/Tirvandor-Players-Handbook-Complete.docx"
        
        self.chapters = [
            '01-introduction.md', '02-character-creation.md', '03-races.md',
            '04-classes.md', '05-backgrounds.md', '06-equipment.md',
            '07-customization.md', '08-custom-subclasses.md', '09-custom-spells.md',
            '10-conditions.md', '11-world-primer.md', '12-legal.md'
        ]
    
    def create_document(self):
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1.0)
            section.left_margin = section.right_margin = Inches(1.0)
        return doc
    
    def set_cell_border(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:color'), '2C3E50')
            tcBorders.append(edge_el)
        
        tcPr.append(tcBorders)
    
    def to_title_case(self, text):
        lowercase_words = {'a', 'an', 'and', 'as', 'at', 'but', 'by', 'for', 
                          'in', 'nor', 'of', 'on', 'or', 'so', 'the', 'to', 'up', 'yet'}
        uppercase_words = {'AC', 'CR', 'DC', 'DM', 'HP', 'NPC', 'PC', 'XP', 
                          'SRD', 'OGL', 'I', 'II', 'III', 'IV', 'V'}
        
        words = text.split()
        result = []
        
        for i, word in enumerate(words):
            clean_word = word.strip('*`')
            
            if clean_word.upper() in uppercase_words:
                result.append(clean_word.upper())
            elif i == 0 or (i > 0 and words[i-1].endswith(':')):
                result.append(clean_word.capitalize())
            elif clean_word.lower() in lowercase_words:
                result.append(clean_word.lower())
            else:
                result.append(clean_word.capitalize())
        
        return ' '.join(result)
    
    def add_cover_page(self, doc):
        print("  📘 Adding cover...")
        
        cover_path = f"{self.images_dir}/tirvandor-cover-players-guide-real.png"
        
        if os.path.exists(cover_path):
            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(cover_path, width=Inches(6.0))
                doc.add_page_break()
                print("    ✅ Cover embedded")
                return True
            except:
                pass
        return False
    
    def extract_full_toc(self):
        """Extract COMPLETE TOC from all markdown files"""
        toc_entries = []
        
        for chapter_file in self.chapters:
            filepath = f"{self.markdown_dir}/{chapter_file}"
            if not os.path.exists(filepath):
                continue
            
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
            
            # Extract ALL H1 and H2 headings
            for line in lines:
                line = line.strip()
                
                # Clean markdown
                clean_line = re.sub(r'\*\*([^\*]+)\*\*', r'\1', line)
                clean_line = re.sub(r'\*([^\*]+)\*', r'\1', clean_line)
                clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)
                clean_line = clean_line.strip()
                
                if line.startswith('# ') and not line.startswith('####'):
                    heading = self.to_title_case(clean_line[2:])
                    toc_entries.append(('h1', heading))
                elif line.startswith('## '):
                    heading = self.to_title_case(clean_line[3:])
                    toc_entries.append(('h2', heading))
        
        return toc_entries
    
    def add_full_toc(self, doc):
        print("  📋 Building FULL table of contents...")
        
        # Title
        heading = doc.add_heading("Contents", 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(139, 0, 0)
            run.font.size = Pt(24)
        
        doc.add_paragraph()
        
        # Get ALL entries
        toc_entries = self.extract_full_toc()
        
        # Add entries
        for level, heading_text in toc_entries:
            para = doc.add_paragraph()
            
            if level == 'h1':
                run = para.add_run(heading_text)
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(139, 0, 0)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.space_before = Pt(8)
            
            elif level == 'h2':
                run = para.add_run("    " + heading_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(47, 79, 79)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_page_break()
        print(f"    ✅ Added {len(toc_entries)} TOC entries")
    
    def clean_markdown(self, text):
        text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^\*]+)\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return re.sub(r'\s+', ' ', text).strip()
    
    def parse_text_formatting(self, text):
        runs_data = []
        parts = re.split(r'(\*\*[^\*]*?\*\*)', text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                runs_data.append(('bold', part[2:-2]))
            else:
                sub_parts = re.split(r'(\*[^\*]+?\*|`[^`]+?`)', part)
                for sub_part in sub_parts:
                    if not sub_part:
                        continue
                    if sub_part.startswith('*') and sub_part.endswith('*') and not sub_part.startswith('**'):
                        runs_data.append(('italic', sub_part[1:-1]))
                    elif sub_part.startswith('`') and sub_part.endswith('`'):
                        runs_data.append(('code', sub_part[1:-1]))
                    else:
                        runs_data.append(('normal', sub_part))
        
        return runs_data
    
    def add_paragraph_formatted(self, doc, text, style=None):
        para = doc.add_paragraph(style=style)
        for run_type, run_text in self.parse_text_formatting(text):
            if not run_text.strip():
                continue
            run = para.add_run(run_text)
            run.font.size = Pt(11)
            if run_type == 'bold':
                run.bold = True
            elif run_type == 'italic':
                run.italic = True
            elif run_type == 'code':
                run.font.name = 'Courier New'
        return para
    
    def create_table(self, doc, table_lines):
        if not table_lines:
            return
        
        rows = []
        for line in table_lines:
            if re.match(r'^\|[\s\-:]+\|', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        max_cols = max(len(row) for row in rows)
        for row in rows:
            while len(row) < max_cols:
                row.append('')
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                try:
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    clean_text = self.clean_markdown(cell_text)
                    run = para.add_run(clean_text)
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.bold = True
                    self.set_cell_border(cell)
                    if r_idx == 0:
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'D4AF37')
                        cell._element.get_or_add_tcPr().append(shading)
                except IndexError:
                    continue
        
        doc.add_paragraph()
    
    def process_chapter(self, doc, filepath):
        if not os.path.exists(filepath):
            return
        
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        
        i = 0
        table_lines = []
        in_table = False
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if not line or line.strip() == '---':
                i += 1
                continue
            
            if line.startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                i += 1
                continue
            
            if in_table:
                self.create_table(doc, table_lines)
                in_table = False
                table_lines = []
            
            if line.startswith('# '):
                text = self.to_title_case(self.clean_markdown(line[2:].strip()))
                h = doc.add_heading(text, 1)
                h.runs[0].font.color.rgb = RGBColor(139, 0, 0)
            elif line.startswith('## '):
                text = self.to_title_case(self.clean_markdown(line[3:].strip()))
                h = doc.add_heading(text, 2)
                h.runs[0].font.color.rgb = RGBColor(47, 79, 79)
            elif line.startswith('### '):
                text = self.to_title_case(self.clean_markdown(line[4:].strip()))
                doc.add_heading(text, 3)
            elif line.startswith('#### '):
                text = self.to_title_case(self.clean_markdown(line[5:].strip()))
                doc.add_heading(text, 4)
            elif line.startswith('- ') or line.startswith('* '):
                self.add_paragraph_formatted(doc, line[2:].strip(), 'List Bullet')
            elif line.strip():
                self.add_paragraph_formatted(doc, line.strip())
            
            i += 1
        
        if in_table and table_lines:
            self.create_table(doc, table_lines)
    
    def build(self):
        print("=" * 70)
        print("BUILDING PHB WITH COMPLETE TOC")
        print("=" * 70)
        
        doc = self.create_document()
        self.add_cover_page(doc)
        self.add_full_toc(doc)
        
        print("\n📚 Processing chapters:")
        for chapter in self.chapters:
            print(f"  • {chapter.replace('.md', ''):.<45}", end='')
            self.process_chapter(doc, f"{self.markdown_dir}/{chapter}")
            print(" ✅")
        
        print("\n💾 Saving...")
        doc.save(self.output_path)
        
        size = os.path.getsize(self.output_path) / (1024 * 1024)
        print("\n" + "=" * 70)
        print("✅ COMPLETE WITH FULL TOC!")
        print("=" * 70)
        print(f"\n📄 {self.output_path}")
        print(f"📊 {size:.2f} MB")
        print("\n✨ All races, classes, and sections in TOC")

if __name__ == "__main__":
    FullTOCBuilder().build()
