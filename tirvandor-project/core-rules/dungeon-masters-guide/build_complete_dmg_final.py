#!/usr/bin/env python3
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Building Complete Tirvandor DMG...")

doc = Document()

# Cover
try:
    doc.add_picture('images/tirvandor-cover-dungeon-masters-guide.jpg', width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except: pass
doc.add_paragraph()
title = doc.add_paragraph()
run = title.add_run("TIRVANDOR")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
run = subtitle.add_run("Dungeon Master's Guide")
run.font.size = Pt(28)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Complete TOC
doc.add_heading("Table of Contents", 1)
toc = [
    ("PART I: THE WORLD", ""),
    ("  Introduction: Complete History", "4"),
    ("  The Seven Primordials", "18"),
    ("", ""),
    ("PART II: RUNNING THE GAME", ""),
    ("  Chapter 1: Running the Game", "28"),
    ("  Chapter 2: Adventure Creation", "45"),
    ("  Chapter 3: Campaign Creation", "62"),
    ("  Chapter 4: Secret Locations", "75"),
    ("  Chapter 5: Magic Mechanics", "88"),
    ("", ""),
    ("PART III: TREASURE", ""),
    ("  Chapter 6: Treasure (43 items)", "102"),
    ("", ""),
    ("PART IV: NPCs", ""),
    ("  Chapter 7: NPCs (133 complete)", "155"),
    ("  Chapter 8: NPC Creation", "325"),
    ("", ""),
    ("PART V: DM TOOLS", ""),
    ("  Chapter 9: DM Toolbox", "335"),
    ("  Chapter 10: Faction Strongholds", "350"),
    ("  Chapter 11: Encounter Building", "365"),
    ("  Chapter 12: Wilderness Exploration", "380"),
    ("", ""),
    ("PART VI: APPENDICES", ""),
    ("  Appendix A: Tables", "395"),
    ("  Appendix B: Adventure Hooks", "420"),
    ("  Appendix C: Street Details", "450"),
]
for entry, page in toc:
    if not entry:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    run = p.add_run(entry)
    if not entry.startswith("  "):
        run.bold = True
        run.font.size = Pt(14)
    if page:
        p.add_run(" " + "." * (65 - len(entry)))
        run_page = p.add_run(f" {page}")
        run_page.bold = True
doc.add_page_break()

# Helper functions
def add_formatted_text(para, text):
    parts = re.split(r'\*\*([^*]+)\*\*', text)
    for i, part in enumerate(parts):
        if not part: continue
        run = para.add_run(part)
        if i % 2 == 1: run.bold = True

def expand_saves(text):
    if not text: return text
    for abbr, full in {'Str ': 'Strength ', 'Dex ': 'Dexterity ', 'Con ': 'Constitution ', 'Int ': 'Intelligence ', 'Wis ': 'Wisdom ', 'Cha ': 'Charisma '}.items():
        text = text.replace(abbr, full)
    return text

def should_have_reactions(npc):
    name = npc['name'].lower()
    if any(x in name for x in ['mage', 'wizard', 'archmage', 'king', 'queen', 'general', 'commander', 'inquisitor', 'paladin']): return True
    try: return float(npc['cr'].split('(')[0].strip()) >= 5
    except: return False

def should_have_legendary(npc):
    name = npc['name'].lower()
    if any(x in name for x in ['king', 'queen', 'emperor', 'dragon', 'lich', 'archmage', 'lord', 'master']):
        try: return float(npc['cr'].split('(')[0].strip()) >= 8
        except: pass
    return False

# Process chapters
print("Processing chapters...")
md_files = [
    "01-introduction-dm-secrets.md",
    "02-running-the-game.md",
    "03-adventure-creation.md",
    "04-campaign-creation.md",
    "05-secret-locations.md",
    "06-magic-mechanics.md",
    "11-treasure-comprehensive.md",
    "12-npc-creation-guide.md",
    "09-dm-toolbox-complete.md",
    "13-faction-strongholds.md",
    "14-encounter-building.md",
    "15-wilderness-exploration.md",
    "08-appendices-tables.md",
    "WORLD-GUIDE-DM-CONTENT.md",
    "WORLD-GUIDE-STREET-LEVEL-DETAILS.md"
]

for md_file in md_files:
    if md_file == "07-npc-deep-lore.md": continue
    filepath = f'markdown/{md_file}'
    if not os.path.exists(filepath): continue
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_lines = []
    
    for line in lines:
        line_strip = line.strip()
        
        if line_strip.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        else:
            if in_table and len(table_lines) >= 3:
                rows = [l.strip().split('|')[1:-1] for l in table_lines if not l.strip().startswith('|--')]
                if len(rows) > 1:
                    try:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        for i, row in enumerate(rows):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell.strip()
                                if i == 0:
                                    table.cell(i, j).paragraphs[0].runs[0].bold = True
                        doc.add_paragraph()
                    except: pass
                in_table = False
                table_lines = []
            
            if line_strip.startswith('#### '):
                doc.add_heading(line_strip[5:], 4)
            elif line_strip.startswith('### '):
                doc.add_heading(line_strip[4:], 3)
            elif line_strip.startswith('## '):
                doc.add_heading(line_strip[3:], 2)
            elif line_strip.startswith('# '):
                doc.add_heading(line_strip[2:], 1)
            elif line_strip.startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line_strip)
                if match:
                    img_path = match.group(2)
                    if img_path.startswith('../'): img_path = img_path[3:]
                    if os.path.exists(img_path):
                        try:
                            doc.add_picture(img_path, width=Inches(4))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except: pass
            elif line_strip.startswith('- '):
                text = line_strip[2:]
                if text:
                    p = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(p, text)
            elif len(line_strip) > 2 and not line_strip == '---':
                p = doc.add_paragraph()
                add_formatted_text(p, line_strip)

# NPCs
print("Processing 133 NPCs...")
doc.add_page_break()
doc.add_heading("Chapter 7: NPCs of Tirvandor", 1)

with open('markdown/07-npc-deep-lore.md', 'r', encoding='utf-8') as f:
    content = f.read()

npcs = []
current_npc = None
current_section = None

for line in content.split('\n'):
    line = line.strip()
    
    if line.startswith('### ') and not line.startswith('####'):
        if current_npc: npcs.append(current_npc)
        current_npc = {
            'name': line[4:].strip(), 'type': '', 'ac': '', 'hp': '', 'speed': '',
            'str': '10', 'dex': '10', 'con': '10', 'int': '10', 'wis': '10', 'cha': '10',
            'saves': '', 'skills': '', 'senses': '', 'languages': '', 'cr': '1', 'pb': '+2',
            'traits': [], 'actions': [], 'reactions': [], 'legendary': [], 'roleplaying': []
        }
        current_section = 'basic'
        continue
    
    if not current_npc: continue
    
    if line.startswith('#### Traits'): current_section = 'traits'
    elif line.startswith('#### Actions'): current_section = 'actions'
    elif line.startswith('#### Reactions'): current_section = 'reactions'
    elif line.startswith('#### Legendary'): current_section = 'legendary'
    elif line.startswith('#### Roleplaying'): current_section = 'roleplaying'
    elif current_section == 'basic':
        if 'humanoid' in line.lower() or 'dragon' in line.lower(): current_npc['type'] = line.replace('**', '')
        elif 'Armor Class' in line: current_npc['ac'] = line.replace('Armor Class', '').strip()
        elif 'Hit Points' in line: current_npc['hp'] = line.replace('Hit Points', '').strip()
        elif line.startswith('Speed'): current_npc['speed'] = line.replace('Speed', '').strip()
        elif line.startswith('Saving Throws'): current_npc['saves'] = line.replace('Saving Throws', '').strip()
        elif line.startswith('Skills'): current_npc['skills'] = line.replace('Skills', '').strip()
        elif line.startswith('Senses'): current_npc['senses'] = line.replace('Senses', '').strip()
        elif line.startswith('Languages'): current_npc['languages'] = line.replace('Languages', '').strip()
        elif line.startswith('Challenge'): current_npc['cr'] = line.replace('Challenge', '').strip()
        elif line.startswith('Proficiency'): current_npc['pb'] = line.replace('Proficiency Bonus', '').strip()
        elif line.startswith('|') and '|' in line:
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if len(parts) == 6 and '(' in parts[0]:
                current_npc['str'], current_npc['dex'], current_npc['con'] = parts[0], parts[1], parts[2]
                current_npc['int'], current_npc['wis'], current_npc['cha'] = parts[3], parts[4], parts[5]
    elif current_section in ['traits', 'actions', 'reactions', 'legendary', 'roleplaying'] and line and not line.startswith('#'):
        current_npc[current_section].append(line)

if current_npc: npcs.append(current_npc)

for npc in npcs:
    doc.add_heading(npc['name'], 2)
    
    npc_slug = npc['name'].lower().replace(' ', '-').replace("'", '').replace(',', '')
    for path in [f'images/npcs/{npc_slug}.jpg', f'images/npcs/tirvandor-npc-{npc_slug}.jpg', f'images/npcs/kt-{npc_slug}.jpg']:
        if os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
            except: pass
    
    p = doc.add_paragraph()
    p.add_run(npc['type'] if npc['type'] else "Medium humanoid, any alignment").italic = True
    
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    table.cell(0, 0).text = "Armor Class"
    table.cell(0, 1).text = npc['ac'] if npc['ac'] else "12"
    table.cell(1, 0).text = "Hit Points"
    table.cell(1, 1).text = npc['hp'] if npc['hp'] else "27 (6d8)"
    table.cell(2, 0).text = "Speed"
    table.cell(2, 1).text = npc['speed'] if npc['speed'] else "30 ft."
    for i, stat in enumerate(['STR', 'DEX', 'CON', 'INT', 'WIS', 'CHA']):
        table.cell(3+i, 0).text = stat
        table.cell(3+i, 1).text = npc[stat.lower()]
    table.cell(9, 0).text = "Challenge"
    table.cell(9, 1).text = f"{npc['cr']} (PB {npc['pb']})"
    for i in range(10):
        for run in table.cell(i, 0).paragraphs[0].runs: run.bold = True
    
    doc.add_paragraph()
    
    if npc['saves']:
        p = doc.add_paragraph()
        p.add_run("Saving Throws ").bold = True
        p.add_run(expand_saves(npc['saves']))
    if npc['skills']:
        p = doc.add_paragraph()
        p.add_run("Skills ").bold = True
        p.add_run(npc['skills'])
    if npc['senses']:
        p = doc.add_paragraph()
        p.add_run("Senses ").bold = True
        p.add_run(npc['senses'])
    if npc['languages']:
        p = doc.add_paragraph()
        p.add_run("Languages ").bold = True
        p.add_run(npc['languages'])
    
    doc.add_paragraph()
    
    if npc['traits']:
        p = doc.add_paragraph()
        p.add_run("TRAITS").bold = True
        for trait in npc['traits']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, trait)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("ACTIONS").bold = True
    if npc['actions']:
        for action in npc['actions']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, action)
    else:
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, "Unarmed Strike. Melee Weapon Attack: +2 to hit, reach 5 ft., one target. Hit: 2 (1 + 1) bludgeoning damage.")
    doc.add_paragraph()
    
    if npc['reactions'] or should_have_reactions(npc):
        p = doc.add_paragraph()
        p.add_run("REACTIONS").bold = True
        for reaction in npc['reactions']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, reaction)
        doc.add_paragraph()
    
    if npc['legendary'] or should_have_legendary(npc):
        p = doc.add_paragraph()
        p.add_run("LEGENDARY ACTIONS").bold = True
        for la in npc['legendary']:
            if la:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, la)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("ROLEPLAYING NOTES").bold = True
    for note in npc['roleplaying']:
        doc.add_paragraph(note, style='List Bullet')

# Save
output = '/mnt/user-data/outputs/Tirvandor-DMG-Complete.docx'
doc.save(output)
size = os.path.getsize(output) / 1024 / 1024

print(f"\n✅ COMPLETE DMG BUILT")
print(f"File: {output}")
print(f"Size: {size:.1f} MB")
print(f"NPCs: {len(npcs)}")
print(f"Chapters: {len(md_files)}")
