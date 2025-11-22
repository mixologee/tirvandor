#!/usr/bin/env python3
"""
Build Complete Tirvandor DMG - Professional Edition
"""

import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("BUILDING TIRVANDOR DMG - PROFESSIONAL COMPLETE EDITION")
print("="*80)

doc = Document()

# COVER PAGE
print("\n✅ Adding cover...")
try:
    doc.add_picture('images/tirvandor-cover-dungeon-masters-guide.jpg', width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"⚠️  Cover image error: {e}")

doc.add_paragraph()

# Title
title = doc.add_paragraph()
run = title.add_run("TIRVANDOR")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
run = subtitle.add_run("Dungeon Master's Guide")
run.font.size = Pt(28)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtext = doc.add_paragraph()
run = subtext.add_run("Complete Edition")
run.font.size = Pt(16)
run.font.italic = True
subtext.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# TABLE OF CONTENTS
print("✅ Creating detailed table of contents...")
doc.add_heading("Table of Contents", 1)

toc_structure = [
    ("Introduction: Secrets Behind the World", 4),
    ("  Complete Timeline (3,700 BS - 1247 CR)", 5),
    ("  The Seven Primordials", 18),
    ("  Divine Mechanics", 22),
    ("Chapter 1: Running the Game", 25),
    ("  Ability Checks and DCs", 26),
    ("  Combat Mechanics", 30),
    ("  Exploration Rules", 35),
    ("Chapter 2: Adventure Creation", 42),
    ("  Step-by-Step Design", 43),
    ("  Encounter Building", 48),
    ("Chapter 3: Campaign Creation", 58),
    ("  Campaign Structure", 59),
    ("  Faction Dynamics", 64),
    ("Chapter 4: Secret Locations", 70),
    ("  Primordial Prison Sites", 71),
    ("  Hidden Places of Power", 78),
    ("Chapter 5: Magic Mechanics", 85),
    ("  Ley Line System", 86),
    ("  Corruption Rules", 91),
    ("  Item Crafting", 96),
    ("Chapter 6: Treasure", 102),
    ("  Comprehensive Tables", 103),
    ("  Magic Items A-Z", 108),
    ("Chapter 7: NPCs of Tirvandor", 115),
    ("  Complete Stat Blocks (133 NPCs)", 116),
    ("Chapter 8: NPC Creation Guide", 195),
    ("Chapter 9: DM Toolbox", 202),
    ("  Creating Content", 203),
    ("  Environmental Hazards", 208),
    ("Chapter 10: Faction Strongholds", 215),
    ("  Stronghold Types & Management", 216),
    ("Chapter 11: Encounter Building", 225),
    ("  Balanced Encounters & Tactics", 226),
    ("Chapter 12: Wilderness Exploration", 235),
    ("  Travel, Survival, Hazards", 236),
    ("Appendix A: Tables & Charts", 245),
    ("  Weather, Encounters, Treasure", 246),
    ("Appendix B: Adventure Hooks", 260),
    ("  Organized by Region", 261),
    ("Appendix C: Street-Level Details", 285),
    ("  City Descriptions", 286),
    ("Index", 325),
]

for entry, page in toc_structure:
    p = doc.add_paragraph()
    run = p.add_run(entry)
    if not entry.startswith("  "):
        run.bold = True
    p.add_run(" " + "." * (75 - len(entry)))
    run_page = p.add_run(f" {page}")
    run_page.bold = True

doc.add_page_break()

# Process all markdown files
print("\n✅ Processing markdown files...")

md_files = [
    "01-introduction-dm-secrets.md",
    "02-running-the-game.md",
    "03-adventure-creation.md",
    "04-campaign-creation.md",
    "05-secret-locations.md",
    "06-magic-mechanics.md",
    "11-treasure-comprehensive.md",
    "07-npc-deep-lore.md",
    "12-npc-creation-guide.md",
    "09-dm-toolbox-complete.md",
    "13-faction-strongholds.md",
    "14-encounter-building.md",
    "15-wilderness-exploration.md",
    "08-appendices-tables.md",
    "WORLD-GUIDE-DM-CONTENT.md",
    "WORLD-GUIDE-STREET-LEVEL-DETAILS.md"
]

def process_markdown_line(line, doc):
    """Process single markdown line"""
    line = line.rstrip()
    
    if not line or line == '---':
        return
    
    # Headings
    if line.startswith('#### '):
        doc.add_heading(line[5:], 4)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 3)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 2)
    elif line.startswith('# '):
        doc.add_heading(line[2:], 1)
    
    # Lists
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif re.match(r'^\d+\.\s', line):
        doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
    
    # Tables (markdown)
    elif line.startswith('|') and not line.startswith('|--'):
        # Skip for now - tables need special handling
        pass
    
    # Bold text
    elif '**' in line:
        p = doc.add_paragraph()
        parts = re.split(r'\*\*([^*]+)\*\*', line)
        for i, part in enumerate(parts):
            if i % 2 == 1:
                p.add_run(part).bold = True
            else:
                if part.strip():
                    p.add_run(part)
    
    # Italic text
    elif '*' in line and '**' not in line:
        p = doc.add_paragraph()
        parts = re.split(r'\*([^*]+)\*', line)
        for i, part in enumerate(parts):
            if i % 2 == 1:
                p.add_run(part).italic = True
            else:
                if part.strip():
                    p.add_run(part)
    
    # Regular text
    else:
        if len(line.strip()) > 3:
            doc.add_paragraph(line)

file_count = 0
for md_file in md_files:
    filepath = os.path.join('markdown', md_file)
    if not os.path.exists(filepath):
        print(f"  ⚠️  Skipping {md_file} - not found")
        continue
    
    print(f"  → {md_file}")
    file_count += 1
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Track if we're in a table
    in_table = False
    table_lines = []
    
    for line in lines:
        # Handle markdown tables
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        else:
            # If we were in table, process it
            if in_table and len(table_lines) >= 3:
                # Create Word table
                rows = [l.strip().split('|')[1:-1] for l in table_lines if not l.strip().startswith('|--')]
                if len(rows) > 1:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    for i, row in enumerate(rows):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = cell.strip()
                    doc.add_paragraph()
                in_table = False
                table_lines = []
            
            # Process regular line
            process_markdown_line(line, doc)
    
    # Add page break after major chapters
    if md_file in ["04-campaign-creation.md", "07-npc-deep-lore.md", "13-faction-strongholds.md", "WORLD-GUIDE-DM-CONTENT.md"]:
        doc.add_page_break()

# Save
output_path = '/mnt/user-data/outputs/Tirvandor-DMG-Professional-Complete.docx'
doc.save(output_path)

size_mb = os.path.getsize(output_path) / 1024 / 1024
para_count = len(doc.paragraphs)
table_count = len(doc.tables)
estimated_pages = para_count // 15

print(f"\n{'='*80}")
print(f"✅ DMG COMPLETE!")
print(f"{'='*80}")
print(f"Output: {output_path}")
print(f"Size: {size_mb:.1f} MB")
print(f"Paragraphs: {para_count:,}")
print(f"Tables: {table_count}")
print(f"Estimated pages: ~{estimated_pages}")
print(f"Files processed: {file_count}/16")
print(f"\n✅ Ready for use!")
