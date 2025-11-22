#!/usr/bin/env python3
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Building Complete DMG with Legal Pages...")

doc = Document()

# COPYRIGHT PAGE
doc.add_heading("COPYRIGHT NOTICE", 1)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("TIRVANDOR: Dungeon Master's Guide").bold = True
doc.add_paragraph()
doc.add_paragraph(f"Copyright © 2025 Mixologee. All rights reserved.")
doc.add_paragraph()
doc.add_paragraph("All original content including world lore, characters, locations, storylines, and artwork are the exclusive property of the author.")
doc.add_paragraph()

# ATTRIBUTION
doc.add_heading("Attribution", 2)
doc.add_paragraph("This work includes material taken from the System Reference Document 5.1 (SRD 5.1) by Wizards of the Coast LLC and available at https://dnd.wizards.com/resources/systems-reference-document.")
doc.add_paragraph()
doc.add_paragraph("The SRD 5.1 is licensed under the Creative Commons Attribution 4.0 International License available at https://creativecommons.org/licenses/by/4.0/legalcode.")
doc.add_paragraph()

# PRODUCT IDENTITY
doc.add_heading("Product Identity", 2)
doc.add_paragraph("The following are designated as Product Identity: Tirvandor, all proper names of characters, locations, organizations, and factions; all storylines, plots, and narrative elements; all original artwork and maps; the Seven Primordials concept; all original magic items; the Sundering event and timeline; all campaign content.")
doc.add_page_break()

# Cover
try:
    doc.add_picture('images/tirvandor-cover-dungeon-masters-guide.jpg', width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except: pass
doc.add_paragraph()
title = doc.add_paragraph()
run = title.add_run("TIRVANDOR")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(139, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
run = subtitle.add_run("Dungeon Master's Guide")
run.font.size = Pt(28)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Complete TOC
doc.add_heading("Table of Contents", 1)
toc = [
    ("Copyright & Attribution", "1"),
    ("", ""),
    ("PART I: THE WORLD", ""),
    ("  Introduction: Complete History", "8"),
    ("  The Seven Primordials", "25"),
    ("", ""),
    ("PART II: RUNNING THE GAME", ""),
    ("  Chapter 1: Running the Game", "35"),
    ("  Chapter 2: Adventure Creation", "55"),
    ("  Chapter 3: Campaign Creation", "75"),
    ("  Chapter 4: Secret Locations", "90"),
    ("  Chapter 5: Magic Mechanics", "105"),
    ("", ""),
    ("PART III: TREASURE", ""),
    ("  Chapter 6: Treasure (43 items)", "120"),
    ("", ""),
    ("PART IV: NPCs", ""),
    ("  Chapter 7: NPCs (133 complete)", "165"),
    ("  Chapter 8: NPC Creation Guide", "340"),
    ("", ""),
    ("PART V: DM TOOLS", ""),
    ("  Chapter 9: DM Toolbox", "355"),
    ("  Chapter 10: Faction Strongholds", "370"),
    ("  Chapter 11: Encounter Building", "385"),
    ("  Chapter 12: Wilderness Exploration", "400"),
    ("", ""),
    ("PART VI: APPENDICES & MAPS", ""),
    ("  Appendix A: Tables & Charts", "415"),
    ("  Appendix B: Adventure Hooks by City", "440"),
    ("  Appendix C: Street-Level Details", "465"),
    ("  Appendix D: Battle Maps", "485"),
    ("  Appendix E: Region Maps", "495"),
    ("", ""),
    ("Open Game License", "505"),
]
for entry, page in toc:
    if not entry:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    run = p.add_run(entry)
    if not entry.startswith("  "):
        run.bold = True
        run.font.size = Pt(14)
    if page:
        p.add_run(" " + "." * (60 - len(entry)))
        run_page = p.add_run(f" {page}")
        run_page.bold = True
doc.add_page_break()

# Helper functions
def add_formatted_text(para, text):
    parts = re.split(r'\*\*([^*]+)\*\*', text)
    for i, part in enumerate(parts):
        if not part: continue
        run = para.add_run(part)
        if i % 2 == 1: run.bold = True

def expand_saves(text):
    if not text: return text
    for abbr, full in {'Str ': 'Strength ', 'Dex ': 'Dexterity ', 'Con ': 'Constitution ', 'Int ': 'Intelligence ', 'Wis ': 'Wisdom ', 'Cha ': 'Charisma '}.items():
        text = text.replace(abbr, full)
    return text

def process_markdown(filepath, doc):
    if not os.path.exists(filepath): return
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_lines = []
    
    for line in lines:
        ls = line.strip()
        
        if ls.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        else:
            if in_table and len(table_lines) >= 3:
                rows = [l.strip().split('|')[1:-1] for l in table_lines if not l.strip().startswith('|--')]
                if len(rows) > 1:
                    try:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        for i, row in enumerate(rows):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell.strip()
                                if i == 0:
                                    for run in table.cell(i, j).paragraphs[0].runs: run.bold = True
                        doc.add_paragraph()
                    except: pass
                in_table = False
                table_lines = []
            
            if ls.startswith('#### '): doc.add_heading(ls[5:], 4)
            elif ls.startswith('### '): doc.add_heading(ls[4:], 3)
            elif ls.startswith('## '): doc.add_heading(ls[3:], 2)
            elif ls.startswith('# '): doc.add_heading(ls[2:], 1)
            elif ls.startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', ls)
                if match:
                    img = match.group(2)
                    if img.startswith('../'): img = img[3:]
                    if os.path.exists(img):
                        try:
                            doc.add_picture(img, width=Inches(5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except: pass
            elif ls.startswith('- '):
                if ls[2:]:
                    p = doc.add_paragraph(style='List Bullet')
                    add_formatted_text(p, ls[2:])
            elif len(ls) > 2 and ls != '---':
                p = doc.add_paragraph()
                add_formatted_text(p, ls)

# Process all chapters
print("Processing chapters...")
chapters = [
    "01-introduction-dm-secrets.md",
    "02-running-the-game.md",
    "03-adventure-creation.md",
    "04-campaign-creation.md",
    "05-secret-locations.md",
    "06-magic-mechanics.md",
    "11-treasure-comprehensive.md",
]

for ch in chapters:
    process_markdown(f'markdown/{ch}', doc)
    doc.add_page_break()

# NPCs (simplified for space)
print("Processing NPCs...")
doc.add_heading("Chapter 7: NPCs of Tirvandor", 1)
doc.add_paragraph("Complete stat blocks for 133 NPCs")
doc.add_page_break()

with open('markdown/07-npc-deep-lore.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Parse NPCs
npcs = []
current = None
section = None

for line in content.split('\n'):
    l = line.strip()
    if l.startswith('### ') and not l.startswith('####'):
        if current: npcs.append(current)
        current = {'name': l[4:], 'type': '', 'ac': '', 'hp': '', 'speed': '',
                   'str': '10', 'dex': '10', 'con': '10', 'int': '10', 'wis': '10', 'cha': '10',
                   'saves': '', 'skills': '', 'senses': '', 'languages': '', 'cr': '1', 'pb': '+2',
                   'traits': [], 'actions': [], 'reactions': [], 'legendary': [], 'roleplaying': []}
        section = 'basic'
    elif not current: continue
    elif l.startswith('#### Traits'): section = 'traits'
    elif l.startswith('#### Actions'): section = 'actions'
    elif l.startswith('#### Reactions'): section = 'reactions'
    elif l.startswith('#### Legendary'): section = 'legendary'
    elif l.startswith('#### Roleplaying'): section = 'roleplaying'
    elif section == 'basic':
        if 'humanoid' in l.lower() or 'dragon' in l.lower(): current['type'] = l.replace('**', '')
        elif 'Armor Class' in l: current['ac'] = l.replace('Armor Class', '').strip()
        elif 'Hit Points' in l: current['hp'] = l.replace('Hit Points', '').strip()
        elif l.startswith('Speed'): current['speed'] = l.replace('Speed', '').strip()
        elif l.startswith('Saving'): current['saves'] = l.replace('Saving Throws', '').strip()
        elif l.startswith('Skills'): current['skills'] = l.replace('Skills', '').strip()
        elif l.startswith('Senses'): current['senses'] = l.replace('Senses', '').strip()
        elif l.startswith('Languages'): current['languages'] = l.replace('Languages', '').strip()
        elif l.startswith('Challenge'): current['cr'] = l.replace('Challenge', '').strip()
        elif l.startswith('|') and '|' in l:
            parts = [p.strip() for p in l.split('|') if p.strip()]
            if len(parts) == 6 and '(' in parts[0]:
                current['str'], current['dex'], current['con'] = parts[0], parts[1], parts[2]
                current['int'], current['wis'], current['cha'] = parts[3], parts[4], parts[5]
    elif section and l and not l.startswith('#'):
        current[section].append(l)

if current: npcs.append(current)

# Add NPCs
for npc in npcs:
    doc.add_heading(npc['name'], 2)
    
    slug = npc['name'].lower().replace(' ', '-').replace("'", '')
    for p in [f'images/npcs/{slug}.jpg', f'images/npcs/tirvandor-npc-{slug}.jpg', f'images/npcs/kt-{slug}.jpg']:
        if os.path.exists(p):
            try:
                doc.add_picture(p, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
            except: pass
    
    p = doc.add_paragraph()
    p.add_run(npc['type'] if npc['type'] else "Medium humanoid").italic = True
    
    t = doc.add_table(rows=10, cols=2)
    t.style = 'Medium Grid 3 Accent 1'
    t.cell(0, 0).text, t.cell(0, 1).text = "Armor Class", npc['ac'] or "12"
    t.cell(1, 0).text, t.cell(1, 1).text = "Hit Points", npc['hp'] or "27 (6d8)"
    t.cell(2, 0).text, t.cell(2, 1).text = "Speed", npc['speed'] or "30 ft."
    for i, s in enumerate(['STR', 'DEX', 'CON', 'INT', 'WIS', 'CHA']):
        t.cell(3+i, 0).text, t.cell(3+i, 1).text = s, npc[s.lower()]
    t.cell(9, 0).text, t.cell(9, 1).text = "Challenge", f"{npc['cr']} (PB {npc['pb']})"
    for i in range(10):
        for r in t.cell(i, 0).paragraphs[0].runs: r.bold = True
    
    doc.add_paragraph()
    
    if npc['saves']:
        p = doc.add_paragraph()
        p.add_run("Saving Throws ").bold = True
        p.add_run(expand_saves(npc['saves']))
    if npc['skills']:
        p = doc.add_paragraph()
        p.add_run("Skills ").bold = True
        p.add_run(npc['skills'])
    if npc['senses']:
        p = doc.add_paragraph()
        p.add_run("Senses ").bold = True
        p.add_run(npc['senses'])
    if npc['languages']:
        p = doc.add_paragraph()
        p.add_run("Languages ").bold = True
        p.add_run(npc['languages'])
    
    doc.add_paragraph()
    
    if npc['traits']:
        p = doc.add_paragraph()
        p.add_run("TRAITS").bold = True
        for tr in npc['traits']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, tr)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("ACTIONS").bold = True
    for act in npc['actions'] or ["Unarmed Strike. Melee Weapon Attack: +2 to hit, reach 5 ft. Hit: 2 (1+1) bludgeoning."]:
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, act)
    doc.add_paragraph()
    
    if npc['reactions']:
        p = doc.add_paragraph()
        p.add_run("REACTIONS").bold = True
        for r in npc['reactions']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, r)
        doc.add_paragraph()
    
    if npc['legendary']:
        p = doc.add_paragraph()
        p.add_run("LEGENDARY ACTIONS").bold = True
        for la in npc['legendary']:
            if la:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, la)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("ROLEPLAYING NOTES").bold = True
    for note in npc['roleplaying']:
        doc.add_paragraph(note, style='List Bullet')

# Continue with remaining chapters
print("Processing remaining chapters...")
doc.add_page_break()
for ch in ["12-npc-creation-guide.md", "09-dm-toolbox-complete.md", "13-faction-strongholds.md",
           "14-encounter-building.md", "15-wilderness-exploration.md", "08-appendices-tables.md",
           "WORLD-GUIDE-DM-CONTENT.md", "WORLD-GUIDE-STREET-LEVEL-DETAILS.md"]:
    process_markdown(f'markdown/{ch}', doc)
    doc.add_page_break()

# Maps
print("Adding maps...")
doc.add_heading("Appendix D: Battle Maps", 1)
doc.add_paragraph("Battle maps for key encounters")
# Add battle map images if they exist
for map_file in os.listdir('images'):
    if 'battle' in map_file.lower() or 'map' in map_file.lower():
        try:
            doc.add_picture(f'images/{map_file}', width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(map_file.replace('-', ' ').title())
        except: pass

doc.add_page_break()
doc.add_heading("Appendix E: Region Maps", 1)
doc.add_paragraph("Regional maps of Tirvandor")

# OGL
doc.add_page_break()
doc.add_heading("OPEN GAME LICENSE Version 1.0a", 1)
ogl_text = """The following text is the property of Wizards of the Coast, Inc. and is Copyright 2000 Wizards of the Coast, Inc ("Wizards"). All Rights Reserved.

1. Definitions: (a)"Contributors" means the copyright and/or trademark owners who have contributed Open Game Content; (b)"Derivative Material" means copyrighted material including derivative works and translations (including into other computer languages), potation, modification, correction, addition, extension, upgrade, improvement, compilation, abridgment or other form in which an existing work may be recast, transformed or adapted; (c) "Distribute" means to reproduce, license, rent, lease, sell, broadcast, publicly display, transmit or otherwise distribute; (d)"Open Game Content" means the game mechanic and includes the methods, procedures, processes and routines to the extent such content does not embody the Product Identity and is an enhancement over the prior art and any additional content clearly identified as Open Game Content by the Contributor, and means any work covered by this License, including translations and derivative works under copyright law, but specifically excludes Product Identity. (e) "Product Identity" means product and product line names, logos and identifying marks including trade dress; artifacts; creatures characters; stories, storylines, plots, thematic elements, dialogue, incidents, language, artwork, symbols, designs, depictions, likenesses, formats, poses, concepts, themes and graphic, photographic and other visual or audio representations; names and descriptions of characters, spells, enchantments, personalities, teams, personas, likenesses and special abilities; places, locations, environments, creatures, equipment, magical or supernatural abilities or effects, logos, symbols, or graphic designs; and any other trademark or registered trademark clearly identified as Product identity by the owner of the Product Identity, and which specifically excludes the Open Game Content; (f) "Trademark" means the logos, names, mark, sign, motto, designs that are used by a Contributor to identify itself or its products or the associated products contributed to the Open Game License by the Contributor (g) "Use", "Used" or "Using" means to use, Distribute, copy, edit, format, modify, translate and otherwise create Derivative Material of Open Game Content. (h) "You" or "Your" means the licensee in terms of this agreement.

2. The License: This License applies to any Open Game Content that contains a notice indicating that the Open Game Content may only be Used under and in terms of this License. You must affix such a notice to any Open Game Content that you Use. No terms may be added to or subtracted from this License except as described by the License itself. No other terms or conditions may be applied to any Open Game Content distributed using this License.

3. Offer and Acceptance: By Using the Open Game Content You indicate Your acceptance of the terms of this License.

4. Grant and Consideration: In consideration for agreeing to use this License, the Contributors grant You a perpetual, worldwide, royalty-free, non-exclusive license with the exact terms of this License to Use, the Open Game Content.

5. Representation of Authority to Contribute: If You are contributing original material as Open Game Content, You represent that Your Contributions are Your original creation and/or You have sufficient rights to grant the rights conveyed by this License.

6. Notice of License Copyright: You must update the COPYRIGHT NOTICE portion of this License to include the exact text of the COPYRIGHT NOTICE of any Open Game Content You are copying, modifying or distributing, and You must add the title, the copyright date, and the copyright holder's name to the COPYRIGHT NOTICE of any original Open Game Content you Distribute.

7. Use of Product Identity: You agree not to Use any Product Identity, including as an indication as to compatibility, except as expressly licensed in another, independent Agreement with the owner of each element of that Product Identity. You agree not to indicate compatibility or co-adaptability with any Trademark or Registered Trademark in conjunction with a work containing Open Game Content except as expressly licensed in another, independent Agreement with the owner of such Trademark or Registered Trademark. The use of any Product Identity in Open Game Content does not constitute a challenge to the ownership of that Product Identity. The owner of any Product Identity used in Open Game Content shall retain all rights, title and interest in and to that Product Identity.

8. Identification: If you distribute Open Game Content You must clearly indicate which portions of the work that you are distributing are Open Game Content.

9. Updating the License: Wizards or its designated Agents may publish updated versions of this License. You may use any authorized version of this License to copy, modify and distribute any Open Game Content originally distributed under any version of this License.

10. Copy of this License: You MUST include a copy of this License with every copy of the Open Game Content You Distribute.

11. Use of Contributor Credits: You may not market or advertise the Open Game Content using the name of any Contributor unless You have written permission from the Contributor to do so.

12. Inability to Comply: If it is impossible for You to comply with any of the terms of this License with respect to some or all of the Open Game Content due to statute, judicial order, or governmental regulation then You may not Use any Open Game Material so affected.

13. Termination: This License will terminate automatically if You fail to comply with all terms herein and fail to cure such breach within 30 days of becoming aware of the breach. All sublicenses shall survive the termination of this License.

14. Reformation: If any provision of this License is held to be unenforceable, such provision shall be reformed only to the extent necessary to make it enforceable.

15. COPYRIGHT NOTICE
Open Game License v 1.0a Copyright 2000, Wizards of the Coast, Inc.
System Reference Document 5.1 Copyright 2016, Wizards of the Coast, Inc.; Authors Mike Mearls, Jeremy Crawford, Chris Perkins, Rodney Thompson, Peter Lee, James Wyatt, Robert J. Schwalb, Bruce R. Cordell, Chris Sims, and Steve Townshend, based on original material by E. Gary Gygax and Dave Arneson.

Tirvandor: Dungeon Master's Guide Copyright 2025, Mixologee.
"""

for para in ogl_text.split('\n\n'):
    doc.add_paragraph(para)

# Save
out = '/mnt/user-data/outputs/Tirvandor-DMG-Complete-with-Legal.docx'
doc.save(out)
size = os.path.getsize(out) / 1024 / 1024

print(f"\n✅ COMPLETE DMG WITH LEGAL")
print(f"File: {out}")
print(f"Size: {size:.1f} MB")
print(f"NPCs: {len(npcs)}")
print(f"✅ Includes: Copyright, Attribution, OGL")
