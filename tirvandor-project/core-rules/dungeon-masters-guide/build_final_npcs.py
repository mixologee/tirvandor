#!/usr/bin/env python3
"""
Build final NPC stat blocks - NO blank line at end
"""

import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("🧙 Building final NPC stat blocks (no blank line at end)...")

# Read NPC file
with open('markdown/07-npc-deep-lore.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Parse NPCs
npcs = []
current_npc = None
current_section = None

lines = content.split('\n')
for line in lines:
    line = line.strip()
    
    if line.startswith('### ') and not line.startswith('####'):
        if current_npc:
            npcs.append(current_npc)
        
        current_npc = {
            'name': line[4:].strip(),
            'type': '',
            'ac': '',
            'hp': '',
            'speed': '',
            'str': '10', 'dex': '10', 'con': '10', 'int': '10', 'wis': '10', 'cha': '10',
            'saves': '',
            'skills': '',
            'senses': '',
            'languages': '',
            'cr': '1',
            'pb': '+2',
            'traits': [],
            'actions': [],
            'reactions': [],
            'legendary': [],
            'roleplaying': []
        }
        current_section = 'basic'
        continue
    
    if not current_npc:
        continue
    
    if line.startswith('#### Traits'):
        current_section = 'traits'
    elif line.startswith('#### Actions'):
        current_section = 'actions'
    elif line.startswith('#### Reactions'):
        current_section = 'reactions'
    elif line.startswith('#### Legendary'):
        current_section = 'legendary'
    elif line.startswith('#### Roleplaying'):
        current_section = 'roleplaying'
    elif line.startswith('---'):
        continue
    elif current_section == 'basic':
        if 'humanoid' in line.lower() or 'dragon' in line.lower():
            current_npc['type'] = line.replace('**', '')
        elif 'Armor Class' in line:
            current_npc['ac'] = line.replace('Armor Class', '').strip()
        elif 'Hit Points' in line:
            current_npc['hp'] = line.replace('Hit Points', '').strip()
        elif line.startswith('Speed'):
            current_npc['speed'] = line.replace('Speed', '').strip()
        elif line.startswith('Saving Throws'):
            current_npc['saves'] = line.replace('Saving Throws', '').strip()
        elif line.startswith('Skills'):
            current_npc['skills'] = line.replace('Skills', '').strip()
        elif line.startswith('Senses'):
            current_npc['senses'] = line.replace('Senses', '').strip()
        elif line.startswith('Languages'):
            current_npc['languages'] = line.replace('Languages', '').strip()
        elif line.startswith('Challenge'):
            current_npc['cr'] = line.replace('Challenge', '').strip()
        elif line.startswith('Proficiency'):
            current_npc['pb'] = line.replace('Proficiency Bonus', '').strip()
        elif line.startswith('|') and '|' in line:
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if len(parts) == 6 and '(' in parts[0]:
                current_npc['str'] = parts[0]
                current_npc['dex'] = parts[1]
                current_npc['con'] = parts[2]
                current_npc['int'] = parts[3]
                current_npc['wis'] = parts[4]
                current_npc['cha'] = parts[5]
    elif current_section == 'traits' and line and not line.startswith('#'):
        current_npc['traits'].append(line)
    elif current_section == 'actions' and line and not line.startswith('#'):
        current_npc['actions'].append(line)
    elif current_section == 'reactions' and line and not line.startswith('#'):
        current_npc['reactions'].append(line)
    elif current_section == 'legendary' and line and not line.startswith('#'):
        current_npc['legendary'].append(line)
    elif current_section == 'roleplaying' and line and not line.startswith('#'):
        current_npc['roleplaying'].append(line)

if current_npc:
    npcs.append(current_npc)

print(f"Parsed {len(npcs)} NPCs")

# Helper functions (same as before)
def should_have_reactions(npc):
    name_lower = npc['name'].lower()
    if any(x in name_lower for x in ['mage', 'wizard', 'archmage', 'king', 'queen', 'general', 'commander', 'inquisitor', 'paladin']):
        return True
    try:
        cr_num = float(npc['cr'].split('(')[0].strip())
        return cr_num >= 5
    except:
        pass
    return False

def should_have_legendary(npc):
    name_lower = npc['name'].lower()
    if any(x in name_lower for x in ['king', 'queen', 'emperor', 'dragon', 'lich', 'archmage', 'lord', 'master']):
        try:
            cr_num = float(npc['cr'].split('(')[0].strip())
            return cr_num >= 8
        except:
            pass
    return False

def add_default_reactions(npc):
    if not npc['reactions'] and should_have_reactions(npc):
        name_lower = npc['name'].lower()
        if 'mage' in name_lower or 'wizard' in name_lower:
            return ["Counterspell. When a creature within 60 feet casts a spell, use reaction to attempt to counter it (as counterspell spell)."]
        elif 'fighter' in name_lower or 'warrior' in name_lower or 'knight' in name_lower:
            return ["Parry. Add +2 to AC against one melee attack that would hit. Must see attacker and be wielding melee weapon."]
        elif 'rogue' in name_lower or 'thief' in name_lower:
            return ["Uncanny Dodge. When an attacker you can see hits you with an attack, halve the attack's damage."]
    return npc['reactions']

def add_default_legendary(npc):
    if not npc['legendary'] and should_have_legendary(npc):
        base = [
            f"{npc['name']} can take 3 legendary actions, choosing from options below. Only one legendary action option can be used at a time and only at the end of another creature's turn. Spent legendary actions are regained at the start of each turn.",
            "",
            "Detect. Makes a Wisdom (Perception) check.",
            "Move. Moves up to half speed without provoking opportunity attacks.",
            "Attack (Costs 2 Actions). Makes one weapon attack."
        ]
        
        name_lower = npc['name'].lower()
        if 'mage' in name_lower or 'wizard' in name_lower:
            base.append("Cantrip (Costs 2 Actions). Casts a cantrip.")
        elif 'king' in name_lower or 'emperor' in name_lower:
            base.append("Inspire (Costs 2 Actions). One ally within 60 ft. gains advantage on next attack roll, save, or ability check.")
        
        return base
    return npc['legendary']

def add_roleplaying_notes(npc):
    if npc['roleplaying']:
        return npc['roleplaying']
    
    name_lower = npc['name'].lower()
    notes = []
    
    if 'king' in name_lower or 'emperor' in name_lower:
        notes = [
            "Speaks with authority and expects obedience",
            "Weighs every decision for political impact",
            "Burdened by responsibility and duty"
        ]
    elif 'mage' in name_lower or 'wizard' in name_lower or 'archmage' in name_lower:
        notes = [
            "Speaks precisely, choosing words carefully",
            "Often distracted by magical curiosities",
            "Values knowledge above all else"
        ]
    elif 'priest' in name_lower or 'cleric' in name_lower:
        notes = [
            "Invokes divine guidance frequently",
            "Compassionate but firm in beliefs",
            "Seeks to guide others to righteousness"
        ]
    elif 'general' in name_lower or 'commander' in name_lower:
        notes = [
            "Direct and tactical in speech",
            "Expects discipline and order",
            "Evaluates situations strategically"
        ]
    elif 'thief' in name_lower or 'rogue' in name_lower:
        notes = [
            "Witty and sarcastic",
            "Always looking for angles and opportunities",
            "Trust is earned, not given"
        ]
    elif 'noble' in name_lower or 'lord' in name_lower or 'lady' in name_lower:
        notes = [
            "Polished manners and courtly speech",
            "Concerned with reputation and status",
            "Plays political games carefully"
        ]
    else:
        notes = [
            "Practical and straightforward",
            "Focused on their duties",
            "Wary of strangers"
        ]
    
    return notes

def expand_save_abbreviations(saves_text):
    if not saves_text:
        return saves_text
    
    replacements = {
        'Str ': 'Strength ',
        'Dex ': 'Dexterity ',
        'Con ': 'Constitution ',
        'Int ': 'Intelligence ',
        'Wis ': 'Wisdom ',
        'Cha ': 'Charisma '
    }
    
    result = saves_text
    for abbr, full in replacements.items():
        result = result.replace(abbr, full)
    
    return result

def add_formatted_text(paragraph, text):
    parts = re.split(r'\*\*([^*]+)\*\*', text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True

# Create document
doc = Document()

doc.add_heading("Final NPC Stat Blocks - Test", 1)
doc.add_paragraph("No blank line at end of entries")
doc.add_paragraph()

# Process first 10 NPCs
for npc in npcs[:10]:
    npc['reactions'] = add_default_reactions(npc)
    npc['legendary'] = add_default_legendary(npc)
    npc['roleplaying'] = add_roleplaying_notes(npc)
    
    # NPC Name
    doc.add_heading(npc['name'], 2)
    
    # Portrait
    npc_slug = npc['name'].lower().replace(' ', '-').replace("'", '').replace(',', '')
    portrait_paths = [
        f'images/npcs/{npc_slug}.jpg',
        f'images/npcs/tirvandor-npc-{npc_slug}.jpg',
        f'images/npcs/kt-{npc_slug}.jpg',
        f'images/npcs/gr-{npc_slug}.jpg',
    ]
    
    for path in portrait_paths:
        if os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
            except:
                pass
    
    # Type line
    p = doc.add_paragraph()
    run = p.add_run(npc['type'] if npc['type'] else "Medium humanoid, any alignment")
    run.italic = True
    
    # ONE COMPLETE STAT TABLE
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    
    table.cell(0, 0).text = "Armor Class"
    table.cell(0, 1).text = npc['ac'] if npc['ac'] else "12"
    
    table.cell(1, 0).text = "Hit Points"
    table.cell(1, 1).text = npc['hp'] if npc['hp'] else "27 (6d8)"
    
    table.cell(2, 0).text = "Speed"
    table.cell(2, 1).text = npc['speed'] if npc['speed'] else "30 ft."
    
    table.cell(3, 0).text = "STR"
    table.cell(3, 1).text = npc['str']
    table.cell(4, 0).text = "DEX"
    table.cell(4, 1).text = npc['dex']
    table.cell(5, 0).text = "CON"
    table.cell(5, 1).text = npc['con']
    table.cell(6, 0).text = "INT"
    table.cell(6, 1).text = npc['int']
    table.cell(7, 0).text = "WIS"
    table.cell(7, 1).text = npc['wis']
    table.cell(8, 0).text = "CHA"
    table.cell(8, 1).text = npc['cha']
    
    table.cell(9, 0).text = "Challenge"
    table.cell(9, 1).text = f"{npc['cr']} (PB {npc['pb']})"
    
    # Bold left column
    for i in range(10):
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.bold = True
    
    doc.add_paragraph()
    
    # Separate lines for Saves, Senses, Languages
    if npc['saves']:
        p = doc.add_paragraph()
        run = p.add_run("Saving Throws ")
        run.bold = True
        p.add_run(expand_save_abbreviations(npc['saves']))
    
    if npc['skills']:
        p = doc.add_paragraph()
        run = p.add_run("Skills ")
        run.bold = True
        p.add_run(npc['skills'])
    
    if npc['senses']:
        p = doc.add_paragraph()
        run = p.add_run("Senses ")
        run.bold = True
        p.add_run(npc['senses'])
    
    if npc['languages']:
        p = doc.add_paragraph()
        run = p.add_run("Languages ")
        run.bold = True
        p.add_run(npc['languages'])
    
    doc.add_paragraph()
    
    # TRAITS
    if npc['traits']:
        p = doc.add_paragraph()
        run = p.add_run("TRAITS")
        run.bold = True
        for trait in npc['traits']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, trait)
        doc.add_paragraph()
    
    # ACTIONS
    p = doc.add_paragraph()
    run = p.add_run("ACTIONS")
    run.bold = True
    if npc['actions']:
        for action in npc['actions']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, action)
    else:
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, "Unarmed Strike. Melee Weapon Attack: +2 to hit, reach 5 ft., one target. Hit: 2 (1 + 1) bludgeoning damage.")
    doc.add_paragraph()
    
    # REACTIONS (only if appropriate)
    if npc['reactions']:
        p = doc.add_paragraph()
        run = p.add_run("REACTIONS")
        run.bold = True
        for reaction in npc['reactions']:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, reaction)
        doc.add_paragraph()
    
    # LEGENDARY ACTIONS (only if appropriate)
    if npc['legendary']:
        p = doc.add_paragraph()
        run = p.add_run("LEGENDARY ACTIONS")
        run.bold = True
        for la in npc['legendary']:
            if la:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, la)
        doc.add_paragraph()
    
    # ROLEPLAYING NOTES - NO BLANK LINE AFTER
    p = doc.add_paragraph()
    run = p.add_run("ROLEPLAYING NOTES")
    run.bold = True
    for note in npc['roleplaying']:
        doc.add_paragraph(note, style='List Bullet')
    # NO doc.add_paragraph() here!
    
    # Separator immediately after
    doc.add_paragraph("─" * 80)

# Save
output = '/mnt/user-data/outputs/Final-NPC-Stat-Blocks.docx'
doc.save(output)

print(f"\n✅ Created final NPC stat blocks")
print(f"File: {output}")
print(f"Size: {os.path.getsize(output) / 1024 / 1024:.1f} MB")
print("\n✅ FIXED: No blank line at end of NPC entries")
print("   Separator goes immediately after roleplaying notes")
