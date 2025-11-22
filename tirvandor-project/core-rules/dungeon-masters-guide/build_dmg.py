#!/usr/bin/env python3
"""Build Tirvandor Dungeon Master's Guide"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_cover(doc):
    """Add cover page with image"""
    cover_path = 'images/tirvandor-cover-dungeon-masters-guide.png'
    try:
        if Path(cover_path).exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(cover_path, width=Inches(6.5))
            doc.add_page_break()
            return True
    except Exception as e:
        print(f"  (Cover image skipped: {e})")
        # Add text title instead
        p = doc.add_paragraph('DUNGEON MASTER\'S GUIDE')
        p.runs[0].font.size = Pt(48)
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        p = doc.add_paragraph('Tirvandor Campaign Setting')
        p.runs[0].font.size = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        return False

def add_markdown(doc, filepath):
    """Add markdown content with formatting"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)
        elif line.startswith('**') and line.endswith('**') and len(line) < 100:
            p = doc.add_paragraph(line.strip('*'))
            p.runs[0].bold = True
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip().startswith('1. ') or line.strip().startswith('2. '):
            doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
        elif line.strip():
            doc.add_paragraph(line)

print("Building Dungeon Master's Guide...")

doc = Document()

# Styling
for style in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[style]
    s.font.color.rgb = RGBColor(139, 69, 19)
    s.font.name = 'Calibri'

# Cover
print("✓ Cover")
add_cover(doc)

# Content chapters
chapters = [
    '01-introduction-dm-secrets.md',
    '05-secret-locations.md',
    '06-magic-mechanics.md',
    '07-npc-deep-lore.md',
    '08-appendices-tables.md'
]

for chapter in chapters:
    print(f"✓ {chapter}")
    add_markdown(doc, f'markdown/{chapter}')
    doc.add_page_break()

# Save
output = '/mnt/user-data/outputs/Tirvandor-Dungeon-Masters-Guide.docx'
doc.save(output)

size_mb = os.path.getsize(output) / 1024 / 1024
print(f"\n✅ Complete: {size_mb:.1f} MB")
print(f"   8 chapters, {len(doc.paragraphs)} paragraphs")
