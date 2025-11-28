#!/usr/bin/env python3
"""
Tirvandor Monster Manual DOCX Builder - FIXED VERSION
Handles all caps, numbered monsters, images, stat blocks
"""

import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

markdown_dir = "markdown"
output_path = "/mnt/user-data/outputs/Tirvandor-Monster-Manual.docx"

seen_images = set()

def clean(text):
    """Remove non-ASCII, markdown, and normalize"""
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('**', '').replace('*', '').replace('`', '').replace('#', '')
    # Remove leading numbers like "1. " or "10. "
    text = re.sub(r'^\d+\.\s+', '', text.strip())
    return text.strip()

def title_case(text):
    """Convert ALL CAPS to Title Case"""
    # Split on spaces and hyphens, capitalize each word
    words = text.replace('-', ' ').split()
    return ' '.join(word.capitalize() for word in words)

def parse_stat_block_table(lines, idx):
    """Parse stat block as table"""
    table_lines = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        if '|---' not in line:
            table_lines.append(line)
        idx += 1
    
    if len(table_lines) < 2:
        return None, idx
    
    data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    
    max_cols = max(len(row) for row in data) if data else 0
    for row in data:
        while len(row) < max_cols:
            row.append('')
    
    return data, idx

def add_image(doc, filename, width=4.5):
    """Add monster image"""
    if not filename or filename in seen_images:
        return False
    
    path = f"/home/claude/tirvandor-project/npc-portraits/{filename}"
    
    if os.path.exists(path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
            p.paragraph_format.space_after = Pt(12)
            seen_images.add(filename)
            return True
        except:
            pass
    return False

def get_monster_image(monster_name):
    """Complete mapping for all 50 monsters"""
    name = monster_name.lower().strip()
    
    image_map = {
        'border bandit': 'tirvandor-monster-border-bandit.jpg',
        'smuggler captain': 'tirvandor-monster-smuggler-captain.jpg',
        'war-scarred veteran': 'tirvandor-monster-war-scarred-veteran.jpg',
        'war scarred veteran': 'tirvandor-monster-war-scarred-veteran.jpg',
        'border wraith': 'tirvandor-monster-border-wraith.jpg',
        'contested land elemental': 'tirvandor-monster-contested-land-elemental.jpg',
        'refugee mob': 'tirvandor-monster-refugee-mob.jpg',
        'scavenger ghoul': 'tirvandor-monster-scavenger-ghoul.jpg',
        'territorial drake': 'tirvandor-monster-territorial-drake.jpg',
        'war beast': 'tirvandor-monster-war-beast.jpg',
        'haunted battlefield': 'tirvandor-monster-haunted-battlefield.jpg',
        'thaldros conscript': 'tirvandor-monster-thaldros-conscript.jpg',
        'thaldros soldier': 'tirvandor-monster-thaldros-soldier.jpg',
        'iron legion enforcer': 'tirvandor-monster-iron-legion-enforcer.jpg',
        'royal guard elite': 'tirvandor-monster-royal-guard-elite.jpg',
        'state inquisitor': 'tirvandor-monster-state-inquisitor.jpg',
        'war mage of thaldros': 'tirvandor-monster-war-mage-of-thaldros.jpg',
        'siege golem': 'tirvandor-monster-siege-golem.jpg',
        "general's champion": 'tirvandor-monster-generals-champion.jpg',
        'generals champion': 'tirvandor-monster-generals-champion.jpg',
        'iron crown knight': 'tirvandor-monster-iron-crown-knight.jpg',
        'lord commander varius': 'tirvandor-monster-lord-commander-varius-military-leader.jpg',
        'aethorian militia': 'tirvandor-monster-aethorian-militia.jpg',
        'resistance fighter': 'tirvandor-monster-resistance-fighter.jpg',
        "people's champion": 'tirvandor-monster-peoples-champion.jpg',
        'peoples champion': 'tirvandor-monster-peoples-champion.jpg',
        'revolutionary mage': 'tirvandor-monster-revolutionary-mage.jpg',
        'chain breaker monk': 'tirvandor-monster-chain-breaker-monk.jpg',
        'guerrilla commander': 'tirvandor-monster-guerrilla-commander.jpg',
        'the liberator': 'tirvandor-monster-the-liberator.jpg',
        'prophesied hero': 'tirvandor-monster-prophesied-hero.jpg',
        'guild recruit': 'tirvandor-monster-guild-recruit.jpg',
        'veteran mercenary': 'tirvandor-monster-veteran-mercenary.jpg',
        'guild enforcer': 'tirvandor-monster-guild-enforcer.jpg',
        'contract killer': 'tirvandor-monster-contract-killer.jpg',
        'iron guild captain': 'tirvandor-monster-iron-guild-captain.jpg',
        "guildmaster's elite": 'tirvandor-monster-guildmasters-elite.jpg',
        'guildmasters elite': 'tirvandor-monster-guildmasters-elite.jpg',
        'garrick ironheart': 'tirvandor-monster-garrick-ironheart-guildmaster.jpg',
        'blessed paladin': 'tirvandor-monster-blessed-paladin.jpg',
        "thandros's justicar": 'tirvandor-monster-thandross-justicar.jpg',
        'thandross justicar': 'tirvandor-monster-thandross-justicar.jpg',
        "aethor's liberator": 'tirvandor-monster-aethors-liberator.jpg',
        'aethors liberator': 'tirvandor-monster-aethors-liberator.jpg',
        "moira's seer": 'tirvandor-monster-moira-seer.jpg',
        'moiras seer': 'tirvandor-monster-moira-seer.jpg',
        "sylvara's wild hunter": 'tirvandor-monster-sylvaras-wild-hunter.jpg',
        'sylvaras wild hunter': 'tirvandor-monster-sylvaras-wild-hunter.jpg',
        "sera's mercy": 'tirvandor-monster-seras-mercy.jpg',
        'seras mercy': 'tirvandor-monster-seras-mercy.jpg',
        "mordain's sentinel": 'tirvandor-monster-mordains-sentinel.jpg',
        'mordains sentinel': 'tirvandor-monster-mordains-sentinel.jpg',
        'fallen champion': 'tirvandor-monster-fallen-champion.jpg',
        'ancient guardian': 'tirvandor-monster-ancient-guardian.jpg',
        'prophecy keeper': 'tirvandor-monster-prophecy-keeper.jpg',
        'forgotten king': 'tirvandor-monster-forgotten-king.jpg',
        'herald of the seven': 'tirvandor-monster-herald-of-the-seven.jpg',
        'corruption spawn': 'tirvandor-monster-corruption-spawn.jpg',
        'war-twisted soldier': 'tirvandor-monster-war-twisted-soldier.jpg',
        'war twisted soldier': 'tirvandor-monster-war-twisted-soldier.jpg',
        'curse bearer': 'tirvandor-monster-curse-bearer.jpg'
    }
    
    return image_map.get(name)

# Initialize
doc = Document()

# Title page
p = doc.add_heading('TIRVANDOR', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(36)
p = doc.add_heading('Monster Manual', 1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# TOC
p = doc.add_heading('TABLE OF CONTENTS', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for entry in ['Introduction', 'Border Creatures', 'Thaldros Military', 
              'Aethoria & Iron Guild', 'Ancient, Ascended & Corrupted']:
    p = doc.add_paragraph(f'• {entry}')
doc.add_page_break()

# Introduction  
p = doc.add_heading('INTRODUCTION', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("""This Monster Manual contains creatures unique to Tirvandor. Each entry includes complete D&D 5e stat blocks, tactical notes, and world integration.

Creatures are organized by region for easy campaign integration.""")
doc.add_page_break()

# Chapters
chapters = [
    ('01-border-creatures.md', 'BORDER CREATURES'),
    ('02-thaldros-military.md', 'THALDROS MILITARY'),
    ('03-aethoria-iron-guild.md', 'AETHORIA & IRON GUILD'),
    ('04-ancient-ascended.md', 'ANCIENT, ASCENDED & CORRUPTED')
]

for filename, title in chapters:
    filepath = os.path.join(markdown_dir, filename)
    if not os.path.exists(filepath):
        print(f"⚠️  {filename} not found, skipping...")
        continue
    
    # Chapter header
    p = doc.add_heading(title, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(139, 0, 0)
    
    # Read file
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = [line.rstrip() for line in f.readlines()]
    
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line:
            idx += 1
            continue
        
        # Monster heading (## or ###)
        if line.startswith('### ') and not '(' in line:
            # Individual monster within subsection
            raw_name = clean(line[4:])
            display_name = title_case(raw_name)
            
            p = doc.add_heading(display_name, 1)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(139, 0, 0)
            
            img_file = get_monster_image(raw_name)
            if img_file and add_image(doc, img_file):
                pass
            
            idx += 1
            continue
        
        if line.startswith('## '):
            raw_name = clean(line[3:])
            
            # Skip section headers
            if '(' in raw_name and ('Creatures' in raw_name or 'Monsters' in raw_name):
                p = doc.add_heading(raw_name, 2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(70, 130, 180)
                idx += 1
                continue
            
            display_name = title_case(raw_name)
            
            # Add heading
            p = doc.add_heading(display_name, 1)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(139, 0, 0)
            
            # Try to add image
            img_file = get_monster_image(raw_name)
            if img_file and add_image(doc, img_file):
                pass  # Image added
            
            idx += 1
            continue
        
        # Section headers
        if line.startswith('### ') or (line.startswith('**') and line.endswith('**')):
            header = clean(line.replace('###', '').replace('**', ''))
            p = doc.add_paragraph(header)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
            idx += 1
            continue
        
        # Tables
        if line.startswith('|'):
            table_data, idx = parse_stat_block_table(lines, idx)
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.rows[i].cells[j]
                        cell.text = clean(cell_text)
                        if i == 0:  # Header row
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.bold = True
                
                doc.add_paragraph()
            continue
        
        # Regular text
        cleaned = clean(line)
        if cleaned:
            doc.add_paragraph(cleaned)
        
        idx += 1
    
    doc.add_page_break()

# Save
doc.save(output_path)
print(f"✅ Monster Manual created: {output_path}")
print(f"📊 Images added: {len(seen_images)}")

# Check file size
import os
size_mb = os.path.getsize(output_path) / (1024 * 1024)
print(f"📏 File size: {size_mb:.1f} MB")
