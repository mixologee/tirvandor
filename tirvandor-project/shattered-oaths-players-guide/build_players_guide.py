#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("🔨 Building Shattered Oaths Players Guide...")

doc = Document()

sections = doc.sections
for section in sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add cover
print("  • Adding cover image...")
cover_path = 'images/cover.jpg'
if os.path.exists(cover_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(6.5))
    doc.add_page_break()
    print("    ✅ Cover added")

chapters = [
    '01-introduction.md', '02-campaign-overview.md', '03-who-you-are.md',
    '04-the-prophecy.md', '05-character-creation.md', '06-prophecy-system.md',
    '07-the-resistance.md', '08-campaign-structure.md', '09-heroic-choices.md',
    '10-starting-campaign.md', '11-character-hooks.md', '12-equipment.md',
    '13-house-rules.md', '14-legal.md'
]

base_dir = 'markdown'
first_chapter = True

for chapter_file in chapters:
    filepath = os.path.join(base_dir, chapter_file)
    if not os.path.exists(filepath):
        print(f"  ⚠️  Skipping {chapter_file}")
        continue
    
    print(f"  • {chapter_file}...", end='')
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            if not first_chapter:
                doc.add_page_break()
            p = doc.add_paragraph(line[2:])
            p.style = 'Heading 1'
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            first_chapter = False
            
        elif line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = 'Heading 3'
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            run = p.runs[0] if p.runs else None
            if run:
                run.font.size = Pt(11)
                
        elif line.strip():
            # Handle bold/italic markdown
            text = line
            p = doc.add_paragraph()
            
            # Handle italics first
            parts = text.split('*')
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(11)
                if i % 2 == 1:  # Odd indexes are italic
                    run.font.italic = True
            
            # If no italics, handle bold
            if '*' not in text and '**' in text:
                p.clear()
                parts = text.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    if i % 2 == 1:
                        run.font.bold = True
            
            pf = p.paragraph_format
            pf.space_after = Pt(6)
        else:
            doc.add_paragraph()
    
    print(" ✅")

output_path = '/mnt/user-data/outputs/Shattered-Oaths-Players-Guide-Rebuilt.docx'
doc.save(output_path)

size = os.path.getsize(output_path)
size_mb = size / (1024 * 1024)
print(f"\n✅ COMPLETE: {size_mb:.2f} MB")
