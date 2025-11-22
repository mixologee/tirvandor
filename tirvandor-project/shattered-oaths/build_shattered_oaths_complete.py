import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

md_dir = "markdown"
images_dir = "images"
cover_path = "images/tirvandor-cover-shattered-oaths-campaign.jpg"

seen_headers = set()
seen_images = set()

def clean(text):
    """Remove markdown formatting and non-ASCII"""
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('**', '').replace('*', '').replace('`', '').replace('#', '')
    return text.strip()

def parse_table(lines, idx):
    """Parse markdown table into data structure"""
    table_lines = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        if '|---' not in line and '|-' not in line:
            table_lines.append(line)
        idx += 1
    if len(table_lines) < 2:
        return None, idx
    data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    max_cols = max(len(row) for row in data) if data else 0
    for row in data:
        while len(row) < max_cols:
            row.append('')
    return data, idx

def add_table_to_doc(doc, data):
    """Add formatted table to document"""
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = clean(cell_data)
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True

def add_image(doc, filename, width=5.5):
    """Add image if not already added"""
    if filename in seen_images:
        return False
    paths = [
        os.path.join(images_dir, filename),
        os.path.join(images_dir, 'npcs', filename),
        os.path.join(images_dir, 'maps', filename),
        os.path.join(images_dir, 'items', filename)
    ]
    for path in paths:
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
            p.paragraph_format.space_after = Pt(12)
            seen_images.add(filename)
            return True
    return False

def match_image(header_text, quest_context=""):
    """Match headers to appropriate images"""
    h = header_text.lower()
    q = quest_context.lower()
    
    # NPC mappings for Shattered Oaths
    npc_mapping = {
        'lord shadows': 'bc-corvus-blackwood-necromancer.jpg',
        'corvus blackwood': 'bc-corvus-blackwood-necromancer.jpg',
        'corvus': 'bc-corvus-blackwood-necromancer.jpg',
        'elara': 'bc-elara-prophet.jpg',
        'helena dawnblade': 'faction-helena-blackstone-iron-legion.jpg',
        'commander helena': 'faction-helena-blackstone-iron-legion.jpg',
        'lyanna thornwood': 'faction-lyanna-thornwood-red-wolf.jpg',
        'aria dawnbringer': 'faction-aria-dawnbringer-paladin.jpg',
        'garrick ironheart': 'faction-garrick-ironheart-guildmaster.jpg',
        'viktor seaworth': 'gr-admiral-viktor-seaworth.jpg',
        'admiral viktor': 'gr-admiral-viktor-seaworth.jpg',
    }
    
    # Map mappings
    map_mapping = {
        'raven': 'ravens-keep-lair.jpg',
        'fortress': 'fortress-siege-battle.jpg',
        'vault': 'sundering-vault.jpg',
        'shrine': 'ancient-shrine.jpg',
        'ruins': 'ancient-ruins.jpg',
        'battle': 'battle-field.jpg',
        'dungeon': 'dungeon-corridor.jpg',
    }
    
    for key, img in npc_mapping.items():
        if key in h:
            return img
    
    for key, img in map_mapping.items():
        if key in h or key in q:
            return img
    
    return None

def process_markdown_file(doc, filepath, quest_context=""):
    """Process markdown file with proper formatting"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Check for tables
        if line.strip().startswith('|'):
            table_data, i = parse_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
            continue
        
        # Headers
        if line.startswith('# ') and not line.startswith('##'):
            text = clean(line[2:])
            doc.add_heading(text, 0)
            img = match_image(text, quest_context)
            if img:
                add_image(doc, img)
        elif line.startswith('## ') and not line.startswith('###'):
            text = clean(line[3:])
            doc.add_heading(text, 1)
            img = match_image(text, quest_context)
            if img:
                add_image(doc, img)
        elif line.startswith('### ') and not line.startswith('####'):
            text = clean(line[4:])
            doc.add_heading(text, 2)
            img = match_image(text, quest_context)
            if img:
                add_image(doc, img)
        elif line.startswith('#### '):
            text = clean(line[5:])
            doc.add_heading(text, 3)
        
        # Read-aloud text (blockquotes)
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            text = text.replace('*', '').replace('_', '')  # Remove emphasis markers
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        
        # Bold paragraphs
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = clean(line)
            p = doc.add_paragraph(text)
            p.runs[0].bold = True
        
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean(line.strip()[2:])
            doc.add_paragraph(text, style='List Bullet')
        
        # Regular paragraphs
        elif line.strip() and not line.startswith('---'):
            text = clean(line)
            doc.add_paragraph(text)
        
        i += 1

# Create document
doc = Document()

# Setup styles
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    try:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style.font.color.rgb = RGBColor(70, 70, 70)
    except:
        pass

# Add cover
if os.path.exists(cover_path):
    for section in doc.sections:
        section.top_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(8.5))
    doc.add_page_break()

# Set margins for content
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Copyright notice
p = doc.add_paragraph('© 2025 Tirvandor Campaign Setting. For personal use only.')
p.runs[0].font.size = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Title page
p = doc.add_paragraph('THE SHATTERED OATHS')
p.runs[0].font.size = Pt(36)
p.runs[0].bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Complete Campaign Guide')
p.runs[0].font.size = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Levels 1-15 | Heroic Fantasy')
p.runs[0].font.size = Pt(14)
p.runs[0].font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Process content
print("Building Shattered Oaths Complete Campaign...")
print("Part 1: Campaign content...")
process_markdown_file(doc, f'{md_dir}/SHATTERED-OATHS-COMPLETE.md')

doc.add_page_break()
print("Part 2: NPCs...")
doc.add_heading('Appendix A: Campaign NPCs', 0)
process_markdown_file(doc, f'{md_dir}/CAMPAIGN-NPCS-SHATTERED-OATHS.md')

doc.add_page_break()
print("Part 3: Items...")
doc.add_heading('Appendix B: Magic Items', 0)
process_markdown_file(doc, f'{md_dir}/CAMPAIGN-ITEMS-SHATTERED-OATHS.md')

doc.add_page_break()
print("Part 4: Monsters...")
doc.add_heading('Appendix C: Monsters', 0)
process_markdown_file(doc, f'{md_dir}/monsters.md')

doc.add_page_break()
print("Part 5: Handouts...")
doc.add_heading('Appendix D: Player Handouts', 0)
process_markdown_file(doc, f'{md_dir}/HANDOUTS-SHATTERED-OATHS.md')

doc.add_page_break()
print("Part 6: Maps...")
doc.add_heading('Appendix E: Battle Maps', 0)
process_markdown_file(doc, f'{md_dir}/SHATTERED-OATHS-MAPS.md')

# Save
output_path = '/mnt/user-data/outputs/Shattered-Oaths-Complete-Campaign.docx'
doc.save(output_path)

size_mb = os.path.getsize(output_path) / 1024 / 1024
print(f"\n✅ Complete: {size_mb:.2f} MB")
print(f"✅ {len(seen_images)} images placed")
