import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean(text):
    return re.sub(r'[^\x00-\x7F]+', '', text).replace('**', '').replace('`', '').replace('#', '').strip()

doc = Document()

# Add cover page
cover_path = 'images/shattered-oaths-cover.jpg'
if os.path.exists(cover_path):
    for section in doc.sections:
        section.top_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(cover_path, width=Inches(8.5))
    doc.add_page_break()

# Set margins for content
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Copyright
cp = doc.add_paragraph()
cr = cp.add_run("© 2024-2025 Tirvandor Campaign Setting.\nFor personal tabletop use only.")
cr.font.size = Pt(8)
cr.italic = True
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Title
doc.add_heading('SHATTERED OATHS', 1).runs[0].font.color.rgb = RGBColor(139, 0, 0)
doc.add_heading('DM Quick Reference Guide', 2).runs[0].font.color.rgb = RGBColor(47, 79, 79)
doc.add_page_break()

# Process the quick reference file
filepath = 'markdown/DM-QUICK-REFERENCE.md'

with open(filepath, 'r', encoding='utf-8') as f:
    lines = f.readlines()

idx = 0
while idx < len(lines):
    line = lines[idx].strip()
    
    if not line or line.startswith('---'):
        idx += 1
        continue
    
    if line.startswith('# '):
        # Skip main title, already added
        if 'DM QUICK REFERENCE' not in line.upper():
            h = doc.add_heading(clean(line[1:]), 1)
            h.runs[0].font.color.rgb = RGBColor(139, 0, 0)
    elif line.startswith('## '):
        h = doc.add_heading(clean(line[2:]), 2)
        h.runs[0].font.color.rgb = RGBColor(47, 79, 79)
    elif line.startswith('### '):
        h = doc.add_heading(clean(line[3:]), 3)
        h.runs[0].font.color.rgb = RGBColor(70, 70, 70)
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(clean(line[2:]), style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
    elif line.startswith('□ '):
        # Checkbox items
        p = doc.add_paragraph(clean(line[2:]), style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
    else:
        # Regular paragraph
        text = clean(line)
        if text:
            p = doc.add_paragraph(text)
            # Bold key terms
            if '**' in line:
                p.clear()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(clean(part))
                    if i % 2 == 1:  # Odd indices are between ** markers
                        run.bold = True
                    run.font.size = Pt(11)
            else:
                p.paragraph_format.space_after = Pt(4)
    
    idx += 1

doc.save("/mnt/user-data/outputs/Shattered-Oaths-DM-Quick-Reference.docx")
print("✅ DM Quick Reference created")
