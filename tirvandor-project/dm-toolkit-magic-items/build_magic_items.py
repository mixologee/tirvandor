import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

script_dir = os.path.dirname(os.path.abspath(__file__))
markdown_dir = os.path.join(script_dir, "markdown")
images_dir = "/home/claude/tirvandor-project/npc-portraits"
cover_path = "/home/claude/tirvandor-project/tirvandor-cover-dm-toolkit-converted.png"
output_path = "/mnt/user-data/outputs/DM-Toolkit-Magic-Items.docx"

def clean_text(text):
    """Clean text but preserve content"""
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    text = text.replace('`', '')
    # Remove box characters but keep content
    box_chars = ['│', '┌', '┐', '└', '┘', '├', '┤', '─', '═', '╔', '╗', '╚', '╝']
    for char in box_chars:
        text = text.replace(char, '')
    return text.strip()

def find_image(filename):
    if not filename:
        return None
    path = os.path.join(images_dir, filename)
    if os.path.exists(path):
        return path
    base = os.path.splitext(filename)[0]
    for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
        path = os.path.join(images_dir, base + ext)
        if os.path.exists(path):
            return path
    return None

def add_cover_and_copyright(doc):
    if os.path.exists(cover_path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(cover_path, width=Inches(8.5), height=Inches(11))
        doc.add_page_break()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    copyright_para = doc.add_paragraph()
    copyright_run = copyright_para.add_run("© 2024-2025 Tirvandor Campaign Setting. All rights reserved.\nFor personal tabletop use only. Not for commercial distribution.")
    copyright_run.font.size = Pt(8)
    copyright_run.italic = True
    copyright_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

print("🔨 Building DM Toolkit - Magic Items (3x3 Grid) with FULL content...")

# Parse items from markdown
with open(os.path.join(markdown_dir, "MAGIC-ITEM-CARDS.md"), 'r') as f:
    lines = f.readlines()

items = []
current_item = {}
in_item = False
in_code_block = False
box_section = 'header'  # header, properties, flavor

for line in lines:
    line_stripped = line.strip()
    
    if line_stripped.startswith('## CARD'):
        if current_item:
            items.append(current_item)
        current_item = {'all_content': []}
        in_item = True
        box_section = 'header'
        # Extract name
        parts = line_stripped.split(':', 1)
        if len(parts) > 1:
            current_item['name'] = parts[1].strip()
    elif in_item and 'Artwork:' in line_stripped:
        artwork = line_stripped.split('`')[1] if '`' in line_stripped else ''
        current_item['artwork'] = artwork
    elif in_item and line_stripped.startswith('```'):
        in_code_block = not in_code_block
    elif in_item and in_code_block:
        # Extract ALL content from box lines
        if '│' in line_stripped:
            content = clean_text(line_stripped)
            if not content:
                continue
            
            # Store all content
            current_item['all_content'].append(content)
            
            # Also parse specific fields for formatting
            if any(x in content for x in ['Wondrous Item', 'Weapon', 'Armor', 'Potion', 'Ring', 'Wand', 'Staff', 'Artifact']):
                if any(r in content for r in ['Uncommon', 'Rare', 'Very Rare', 'Legendary', 'Artifact', 'Beyond Legendary']):
                    current_item['type_rarity'] = content
            elif 'Attunement' in content or 'No Attunement' in content or 'DM Control' in content or content.startswith('('):
                current_item['attunement'] = content

if current_item:
    items.append(current_item)

print(f"Found {len(items)} magic items")

# Create document
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

add_cover_and_copyright(doc)

doc.add_heading('Magic Item Cards', 1).runs[0].font.color.rgb = RGBColor(139, 69, 19)
doc.add_paragraph("Complete Collection - 3x3 Grid Format")
doc.add_page_break()

# Create 3x3 grids (9 items per page)
images_added = 0
for page_start in range(0, len(items), 9):
    page_items = items[page_start:page_start + 9]
    
    # Create 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Set row heights
    for row in table.rows:
        row.height = Inches(3.0)
    
    # Fill cells
    for i, item in enumerate(page_items):
        row_idx = i // 3
        col_idx = i % 3
        cell = table.rows[row_idx].cells[col_idx]
        
        # Clear default paragraph
        cell.paragraphs[0].text = ''
        
        # Add image (smaller to make room for all content)
        if 'artwork' in item and item['artwork']:
            image_path = find_image(item['artwork'])
            if image_path:
                try:
                    p = cell.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(image_path, width=Inches(1.5))
                    p.paragraph_format.space_after = Pt(2)
                    images_added += 1
                except Exception as e:
                    print(f"  ⚠️ Failed to add image {item['artwork']}: {e}")
        
        # Add item name (from parsed field)
        p_name = cell.add_paragraph()
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_run = p_name.add_run(item.get('name', 'Unknown Item'))
        name_run.bold = True
        name_run.font.size = Pt(8)
        p_name.paragraph_format.space_after = Pt(2)
        
        # Add type/rarity
        if 'type_rarity' in item:
            p_type = cell.add_paragraph()
            p_type.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            type_run = p_type.add_run(item['type_rarity'])
            type_run.font.size = Pt(6)
            p_type.paragraph_format.space_after = Pt(1)
        
        # Add attunement
        if 'attunement' in item:
            p_att = cell.add_paragraph()
            p_att.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            att_run = p_att.add_run(item['attunement'])
            att_run.font.size = Pt(6)
            att_run.italic = True
            p_att.paragraph_format.space_after = Pt(2)
        
        # Add ALL other content from the box
        if 'all_content' in item:
            skip_count = 0
            for content in item['all_content']:
                # Skip name line (already added)
                if content == item.get('name', ''):
                    continue
                # Skip type/rarity (already added)
                if content == item.get('type_rarity', ''):
                    continue
                # Skip attunement (already added)
                if content == item.get('attunement', ''):
                    continue
                # Skip empty lines
                if not content:
                    continue
                
                # Add the content line
                p_content = cell.add_paragraph()
                
                # Format based on content type
                if content.startswith('•') or content.startswith('+'):
                    # Property bullet point
                    p_content.text = content
                    p_content.runs[0].font.size = Pt(6)
                    p_content.paragraph_format.left_indent = Pt(4)
                elif content.startswith('"') or content.endswith('"'):
                    # Flavor text quote
                    p_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p_content.add_run(content)
                    run.font.size = Pt(6)
                    run.italic = True
                elif content.startswith('Campaign:') or content.startswith('NPC:'):
                    # Attribution
                    p_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p_content.add_run(content)
                    run.font.size = Pt(5)
                    run.italic = True
                else:
                    # Other text
                    p_content.text = content
                    p_content.runs[0].font.size = Pt(6)
                
                p_content.paragraph_format.space_after = Pt(1)
    
    # Add page break after each grid (except last)
    if page_start + 9 < len(items):
        doc.add_page_break()

doc.save(output_path)

print(f"\n✅ Complete: {output_path}")
print(f"   Total items: {len(items)}")
print(f"   Total images: {images_added}/{len(items)}")
print(f"   Pages: {(len(items) + 8) // 9} (9 items per page)")
print(f"   Source: {markdown_dir}/")
   print(f"   Content: ALL box content included")
print(f"   Source: {markdown_dir}/MAGIC-ITEM-CARDS.md")
